"""
Form Engine - Core document generation logic
Adapted from modul_create_temple for FastAPI service
"""

import os
import json
import logging
import subprocess
import datetime
import hashlib
from typing import Dict, Any, Optional, List
from pathlib import Path

from .config import get_settings

# Setup Logging
logger = logging.getLogger(__name__)

# =============================================================================
# GLOBAL RULES - HELPER FUNCTIONS
# =============================================================================

# Checkbox constants - Always use [x]/[ ] instead of unicode checkboxes
CHECKBOX_CHECKED = "[x]"
CHECKBOX_UNCHECKED = "[ ]"


def fill_cell_text(cell, text: str, align=None, bold: bool = False):
    """
    Fill text into a Word table cell with standard formatting.

    Args:
        cell: Word table cell object
        text: Text to fill
        align: Alignment (None = keep original, or use WD_ALIGN_PARAGRAPH.*)
        bold: True = bold, False = normal (default: False to avoid inheriting template bold)
    """
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        WD_ALIGN_PARAGRAPH = None

    # Clear old runs
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""

    # Create new run with text and format
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.text = text
        for run in p.runs:
            run.bold = bold
            run.font.name = 'Times New Roman'

        if align and WD_ALIGN_PARAGRAPH:
            p.alignment = align


def set_left_align_for_lists(doc, var_names: List[str] = None):
    """
    Left-align paragraphs containing bullet lists.

    Auto-detects and left-aligns lines that:
    - Contain variable names from var_names
    - Start with '-' (dash)
    - Start with '•' (bullet)
    """
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return

    def is_list_paragraph(text):
        text = text.strip()
        if var_names:
            for var in var_names:
                if var in text:
                    return True
        if text.startswith('-') or text.startswith('•'):
            return True
        return False

    # Process body paragraphs
    for p in doc.paragraphs:
        if is_list_paragraph(p.text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Process table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if is_list_paragraph(p.text):
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def get_date_line(city: str = "Nam Dinh", day: str = None, month: str = None, year: str = None) -> str:
    """
    Create date line with non-breaking spaces to prevent line wrapping.

    Returns:
        String like: "Nam Dinh, ngay 20 thang 01 nam 2024" with non-breaking spaces
    """
    if day is None or month is None or year is None:
        now = datetime.datetime.now()
        day = str(now.day)
        month = str(now.month)
        year = str(now.year)

    return f"{city},\u00A0ngay\u00A0{day}\u00A0thang\u00A0{month}\u00A0nam\u00A0{year}"


def calculate_sha256(file_path: str) -> str:
    """Calculate SHA256 hash of a file"""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


class FormEngine:
    """Main document generation engine"""

    def __init__(self, template_dir: str = None, output_dir: str = None, log_dir: str = None):
        settings = get_settings()
        self.template_dir = template_dir or settings.template_dir
        self.output_dir = output_dir or settings.output_dir
        self.log_dir = log_dir or settings.log_dir
        self.base_url = settings.base_url

        # Create today's output directory
        self.today_dir = os.path.join(self.output_dir, datetime.datetime.now().strftime("%Y-%m-%d"))
        os.makedirs(self.today_dir, exist_ok=True)
        os.makedirs(self.log_dir, exist_ok=True)

    def _get_output_path(self, base_name: str, ext: str) -> str:
        """Generate unique output file path with timestamp"""
        timestamp = datetime.datetime.now().strftime("%H%M%S")
        return os.path.join(self.today_dir, f"{base_name}_{timestamp}.{ext}")

    def _replace_text_in_element(self, element, context: Dict[str, Any]):
        """Replace template variables in a paragraph or cell element"""
        if not element.text:
            return

        for key, val in context.items():
            # Support multiple patterns: {{key}}, {{ key }}, {{key }}, {{ key}}
            patterns = [
                f"{{{{ {key} }}}}",
                f"{{{{{key}}}}}",
                f"{{{{ {key}}}}}",
                f"{{{{{key} }}}}"
            ]

            val_str = str(val) if val is not None else ""

            for pat in patterns:
                if pat in element.text:
                    # Handle newlines by adding breaks
                    if '\n' in val_str:
                        parts = val_str.split('\n')
                        # Clear old runs
                        for run in element.runs:
                            run.text = ""

                        # Add parts with breaks
                        if element.runs:
                            r = element.runs[0]
                            r.text = parts[0]
                            for part in parts[1:]:
                                r.add_break()
                                r.add_text(part)
                        else:
                            r = element.add_run(parts[0])
                            for part in parts[1:]:
                                r.add_break()
                                r.add_text(part)
                    else:
                        # Normal replace
                        new_text = element.text.replace(pat, val_str)
                        for run in element.runs:
                            run.text = ""
                        if element.runs:
                            element.runs[0].text = new_text
                        else:
                            element.add_run(new_text)

    def get_available_templates(self) -> List[Dict[str, Any]]:
        """List all available templates"""
        templates = []
        template_path = Path(self.template_dir)

        if not template_path.exists():
            return templates

        for f in template_path.glob("*.docx"):
            if f.name.startswith("~$"):  # Skip temp files
                continue
            templates.append({
                "name": f.name,
                "path": str(f),
                "size": f.stat().st_size,
                "modified": datetime.datetime.fromtimestamp(f.stat().st_mtime).isoformat()
            })

        return sorted(templates, key=lambda x: x["name"])

    def get_template_info(self, template_name: str) -> Optional[Dict[str, Any]]:
        """Get info about a specific template"""
        template_path = Path(self.template_dir) / template_name

        if not template_path.exists():
            return None

        return {
            "name": template_name,
            "path": str(template_path),
            "size": template_path.stat().st_size,
            "modified": datetime.datetime.fromtimestamp(template_path.stat().st_mtime).isoformat()
        }

    def render(
        self,
        template_name: str,
        context: Dict[str, Any],
        user_id: str = "system",
        proposal_id: str = None
    ) -> Dict[str, Any]:
        """
        Render a template with provided context data.

        Args:
            template_name: Name of the template file (e.g., "1b.docx")
            context: Dictionary of variables to replace in template
            user_id: ID of user generating the document
            proposal_id: Optional proposal ID for tracking

        Returns:
            Dictionary with paths to generated DOCX and PDF files
        """
        from docx import Document

        template_path = os.path.join(self.template_dir, template_name)
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template '{template_name}' not found")

        base_name = template_name.replace(".docx", "")
        docx_output_path = self._get_output_path(base_name, "docx")

        # Load Document
        doc = Document(template_path)

        # 1. Replace in body paragraphs
        for p in doc.paragraphs:
            self._replace_text_in_element(p, context)

        # 2. Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_text_in_element(p, context)

        # 3. Replace in headers
        for section in doc.sections:
            for p in section.header.paragraphs:
                self._replace_text_in_element(p, context)

        # 4. Replace in footers
        for section in doc.sections:
            for p in section.footer.paragraphs:
                self._replace_text_in_element(p, context)

        # 5. Apply list alignment
        set_left_align_for_lists(doc)

        # Save DOCX
        doc.save(docx_output_path)
        logger.info(f"Generated DOCX: {docx_output_path}")

        # Convert to PDF
        pdf_output_path = None
        try:
            cmd_check = ["soffice", "--version"]
            subprocess.run(cmd_check, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=5)

            cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", self.today_dir, docx_output_path]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)

            expected_pdf = docx_output_path.replace(".docx", ".pdf")
            if os.path.exists(expected_pdf):
                pdf_output_path = expected_pdf
                logger.info(f"Generated PDF: {pdf_output_path}")
        except subprocess.TimeoutExpired:
            logger.warning("PDF conversion timed out")
        except Exception as e:
            logger.warning(f"PDF conversion failed: {e}")

        # Calculate hashes
        sha256_docx = calculate_sha256(docx_output_path)
        sha256_pdf = calculate_sha256(pdf_output_path) if pdf_output_path else None

        # Build relative paths for URLs and API responses
        relative_path = os.path.relpath(docx_output_path, self.output_dir)
        relative_pdf_path = relative_path.replace('.docx', '.pdf') if pdf_output_path else None
        docx_url = f"{self.base_url}/files/{relative_path}"
        pdf_url = f"{self.base_url}/files/{relative_pdf_path}" if pdf_output_path else None

        result = {
            "docx_path": relative_path,  # Return relative path for frontend
            "pdf_path": relative_pdf_path,  # Return relative path for frontend
            "docx_url": docx_url,
            "pdf_url": pdf_url,
            "template": template_name,
            "timestamp": datetime.datetime.now().isoformat(),
            "user_id": user_id,
            "proposal_id": proposal_id,
            "sha256_docx": sha256_docx,
            "sha256_pdf": sha256_pdf
        }

        # Audit Log
        audit_path = os.path.join(self.log_dir, "audit.jsonl")
        with open(audit_path, "a", encoding='utf-8') as f:
            f.write(json.dumps(result, ensure_ascii=False) + "\n")

        return result
