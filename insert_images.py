#!/usr/bin/env python3
"""
insert_images.py - Chèn ảnh vào BAO_CAO_DATN_v7_fixed.docx
Chèn ảnh TRƯỚC mỗi caption "Hình X.X."
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from lxml import etree

DOCX_IN  = '/mnt/dulieu/DoAn/BAO_CAO_DATN_v7_fixed.docx'
DOCX_OUT = '/mnt/dulieu/DoAn/BAO_CAO_DATN_v8_with_images.docx'

SS  = '/mnt/dulieu/DoAn/baocao2/screenshots'
SS2 = '/mnt/dulieu/DoAn/baocao2/diagrams/screenshots'
DG  = '/mnt/dulieu/DoAn/diagrams'

# Map: text caption → đường dẫn ảnh
IMAGE_MAP = {
    'Hình 1.1. Sơ đồ Use Case tổng quát':
        f'{DG}/use_case_diagram.png',

    'Hình 3.1. Sơ đồ kiến trúc tổng thể hệ thống':
        f'{DG}/architecture_diagram.png',

    'Hình 3.2. Sơ đồ thực thể quan hệ (ERD)':
        f'{DG}/erd_diagram.png',

    'Hình 3.3. Sơ đồ máy trạng thái workflow':
        f'{DG}/state_machine_diagram.png',

    'Hình 4.1. Giao diện đăng nhập':
        f'{SS}/01_login_page.png',

    'Hình 4.2. Bảng điều khiển vai trò Giảng viên':
        f'{SS2}/08_gv_dashboard_full.png',

    'Hình 4.3. Bảng điều khiển vai trò Quản lý Khoa':
        f'{SS2}/faculty-dashboard-charts-working.png',

    'Hình 4.4. Bảng điều khiển vai trò Ban Giám hiệu':
        f'{SS2}/bgh-dashboard-charts-working.png',

    'Hình 4.5. Danh sách đề tài NCKH':
        f'{SS2}/09_gv_proposals_list.png',

    'Hình 4.6. Chi tiết đề tài ở trạng thái Nháp':
        f'{SS2}/10_gv_proposal_detail_draft.png',

    'Hình 4.7. Biểu mẫu tạo đề tài mới':
        f'{SS2}/11_gv_create_proposal_form.png',

    'Hình 4.8. Đề tài đã được phê duyệt':
        f'{SS2}/04-proposal-approved.png',

    'Hình 4.9. Phân bổ Hội đồng cho đề tài':
        f'{SS2}/council-dashboard-success.png',

    'Hình 4.11. Quản lý tài khoản người dùng':
        f'{SS2}/13_admin_user_management.png',

    'Hình 4.12. Nhật ký kiểm toán hệ thống':
        f'{SS2}/14_admin_audit_log.png',
}

def add_image_before_paragraph(doc, para_idx, image_path, width_inches=5.5):
    """Chèn paragraph chứa ảnh vào trước para_idx."""
    para = doc.paragraphs[para_idx]
    # Tạo paragraph mới
    new_para = deepcopy(para._element)
    # Clear content của new_para
    for child in list(new_para):
        new_para.remove(child)

    # Tạo run chứa ảnh bằng python-docx
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    # Di chuyển img_para lên trước caption
    para._element.addprevious(img_para._element)
    return img_para


def fix_caption_alignment(doc):
    """Căn giữa tất cả caption Hình."""
    for para in doc.paragraphs:
        if para.text.strip().startswith('Hình '):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # In nghiêng caption
            for run in para.runs:
                run.italic = True


def add_missing_figure_410(doc):
    """Thêm Hình 4.10 còn thiếu (giữa 4.9 và 4.11)."""
    # Tìm Hình 4.11
    for i, para in enumerate(doc.paragraphs):
        if 'Hình 4.11.' in para.text:
            # Thêm placeholder Hình 4.10 trước 4.11
            new_caption = doc.add_paragraph('Hình 4.10. Quản lý Hội đồng đánh giá')
            new_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in new_caption.runs:
                run.italic = True
            para._element.addprevious(new_caption._element)

            # Thêm ảnh trước caption 4.10
            img_path = f'{SS2}/council-dashboard-success.png'
            if os.path.exists(img_path):
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                new_caption._element.addprevious(img_para._element)
            break


def fix_bảng_backend(doc):
    """Bảng so sánh PHP/Java/NestJS bị cắt cụt ở mục 2.2 — log ra để check."""
    for i, para in enumerate(doc.paragraphs):
        if '--------' in para.text:
            print(f'  ⚠️  Bảng bị cắt tại đoạn [{i}]: {para.text[:80]}')


def main():
    print(f'📖 Đọc file: {DOCX_IN}')
    doc = Document(DOCX_IN)

    inserted = 0
    missing = []

    # Duyệt ngược để index không bị lệch khi chèn
    paras = doc.paragraphs
    for i in range(len(paras) - 1, -1, -1):
        text = paras[i].text.strip()
        for caption, img_path in IMAGE_MAP.items():
            if text == caption:
                if os.path.exists(img_path):
                    print(f'  ✅ Chèn ảnh trước [{i}] {caption[:50]}')
                    add_image_before_paragraph(doc, i, img_path)
                    inserted += 1
                else:
                    print(f'  ❌ Thiếu ảnh: {img_path}')
                    missing.append((caption, img_path))
                break

    # Thêm Hình 4.10 còn thiếu
    print('\n📝 Thêm Hình 4.10...')
    add_missing_figure_410(doc)

    # Căn giữa + in nghiêng tất cả caption
    print('🎨 Format captions...')
    fix_caption_alignment(doc)

    # Check bảng bị cắt
    print('\n🔍 Check bảng bị cắt:')
    fix_bảng_backend(doc)

    print(f'\n💾 Lưu file: {DOCX_OUT}')
    doc.save(DOCX_OUT)

    print(f'\n✅ Xong! Đã chèn {inserted} ảnh.')
    if missing:
        print(f'⚠️  {len(missing)} ảnh thiếu:')
        for cap, path in missing:
            print(f'   - {cap}: {path}')


if __name__ == '__main__':
    main()
