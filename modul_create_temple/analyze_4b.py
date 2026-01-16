from docx import Document
import os

def analyze_table_4b():
    path = "form_engine/templates/4b.docx"
    if not os.path.exists(path):
        print("❌ File not found.")
        return

    doc = Document(path)
    if len(doc.tables) == 0:
        print("❌ No tables found inside 4b.docx")
        return

    print(f"✅ Tổng số bảng trong file: {len(doc.tables)}")
    
    for idx, table in enumerate(doc.tables):
        print(f"\n--- BẢNG THỨ {idx + 1} (Index {idx}) ---")
        print(f"Kích thước: {len(table.rows)} dòng x {len(table.columns)} cột")
        
        # Đọc dòng đầu tiên (Header)
        if len(table.rows) > 0:
            headers = []
            for cell in table.rows[0].cells:
                text = cell.text.strip().replace('\n', ' ')[:50] # Lấy 50 ký tự đầu
                headers.append(text)
            print(f"Header: {headers}")
        else:
            print("Bảng rỗng")

if __name__ == "__main__":
    analyze_table_4b()
