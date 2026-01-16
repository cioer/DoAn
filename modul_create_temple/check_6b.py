from docx import Document
import re

def check_6b_tags():
    doc = Document("form_engine/templates/6b.docx")
    text = ""
    for p in doc.paragraphs: text += p.text
    # Scan Header
    for section in doc.sections:
        for p in section.header.paragraphs: text += p.text
        
    tags = re.findall(r'\{\{\s*(\w+)\s*\}\}', text)
    print(f"TAGS IN 6B: {set(tags)}")

if __name__ == "__main__":
    check_6b_tags()
