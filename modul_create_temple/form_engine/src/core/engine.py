import os
import json
import logging
import subprocess
import datetime
from typing import Dict, Any, Optional, List

# Setup Logging
logging.basicConfig(
    filename='form_engine/logs/engine.log',
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# =============================================================================
# QUY TẮC CHUNG (GLOBAL RULES) - HELPER FUNCTIONS
# =============================================================================

# Checkbox constants - LUÔN dùng [x]/[ ] thay vì ☑/☐
CHECKBOX_CHECKED = "[x]"
CHECKBOX_UNCHECKED = "[ ]"


def fill_cell_text(cell, text: str, align=None, bold=False):
    """
    QUY TẮC CHUNG: Điền text vào cell với format chuẩn.

    Args:
        cell: Word table cell object
        text: Text để điền
        align: Căn lề (None = giữ nguyên, hoặc dùng WD_ALIGN_PARAGRAPH.*)
        bold: True = in đậm, False = bình thường (mặc định: False)

    Note:
        - Mặc định bold=False để tránh kế thừa format in đậm từ template
    """
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        WD_ALIGN_PARAGRAPH = None

    # Xóa tất cả runs cũ
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""

    # Tạo run mới với text và format
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.text = text
        for run in p.runs:
            run.bold = bold
            run.font.name = 'Times New Roman'

        if align and WD_ALIGN_PARAGRAPH:
            p.alignment = align


def set_left_align_for_lists(doc, var_names: List[str] = None):
    """
    QUY TẮC CHUNG: Căn trái các paragraph chứa danh sách gạch đầu dòng.

    Args:
        doc: Word Document object
        var_names: List tên biến chứa danh sách (optional)
                  Nếu không cung cấp, tự động nhận diện dòng bắt đầu bằng '-' hoặc '•'

    Note:
        Tự động nhận diện và căn trái các dòng:
        - Chứa tên biến trong var_names
        - Bắt đầu bằng '-' (gạch ngang)
        - Bắt đầu bằng '•' (bullet point)
    """
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return

    def is_list_paragraph(text):
        """Kiểm tra paragraph có phải là danh sách không"""
        text = text.strip()
        if var_names:
            for var in var_names:
                if var in text:
                    return True
        # Kiểm tra ký tự đặc biệt
        if text.startswith('-') or text.startswith('•'):
            return True
        return False

    # Xử lý paragraphs trong body
    for p in doc.paragraphs:
        if is_list_paragraph(p.text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Xử lý paragraphs trong tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if is_list_paragraph(p.text):
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def get_date_line(city: str = "Nam Định", day: str = None, month: str = None, year: str = None) -> str:
    """
    QUY TẮC CHUNG: Tạo dòng ngày tháng không bị ngắt dòng.

    Args:
        city: Tên thành phố
        day, month, year: Giá trị ngày tháng (nếu None, dùng ngày hiện tại)

    Returns:
        Chuỗi dạng: "Nam Định, ngày 20 tháng 01 năm 2024" với non-breaking spaces

    Note:
        Dùng \u00A0 (non-breaking space) thay vì space thường
        để tránh dòng ngày tháng bị ngắt đôi khi xuống trang.
    """
    if day is None or month is None or year is None:
        now = datetime.datetime.now()
        day = str(now.day)
        month = str(now.month)
        year = str(now.year)

    return f"{city},\u00A0ngày\u00A0{day}\u00A0tháng\u00A0{month}\u00A0năm\u00A0{year}"

class FormEngine:
    def __init__(self, template_dir="form_engine/templates", output_dir="form_engine/output"):
        self.template_dir = template_dir
        self.output_dir = output_dir
        self.today_dir = os.path.join(self.output_dir, datetime.datetime.now().strftime("%Y-%m-%d"))
        os.makedirs(self.today_dir, exist_ok=True)

    def _get_output_path(self, base_name: str, ext: str) -> str:
        timestamp = datetime.datetime.now().strftime("%H%M%S")
        return os.path.join(self.today_dir, f"{base_name}_{timestamp}.{ext}")

    def replace_text_in_element(self, element, context):
        """Helper to replace text in paragraph or cell"""
        if not element.text:
            return

        for key, val in context.items():
            # Check for multiple patterns: {{key}}, {{ key }}, {{key }}
            patterns = [
                f"{{{{ {key} }}}}",
                f"{{{{{key}}}}}",
                f"{{{{ {key}}}}}",
                f"{{{{{key} }}}}"
            ]
            
            val_str = str(val) if val is not None else ""
            
            for pat in patterns:
                if pat in element.text:
                    # SMART REPLACE STRATEGY:
                    # If the pattern is split across runs, simple replace won't work on individual runs.
                    # We modify the whole paragraph text.
                    # Warning: This resets formatting of the whole paragraph to the style of the first run.
                    # But ensures data is filled.
                    
                    # Support Newline \n by adding breaks
                    if '\n' in val_str:
                        parts = val_str.split('\n')
                        # Clear old runs
                        for run in element.runs:
                            run.text = ""
                        
                        # Add parts with breaks
                        if element.runs:
                            r = element.runs[0]
                            r.text = parts[0]
                            for part in parts[1:]:
                                r.add_break()
                                r.add_text(part)
                        else:
                            r = element.add_run(parts[0])
                            for part in parts[1:]:
                                r.add_break()
                                r.add_text(part)
                    else:
                        # Normal replace
                        new_text = element.text.replace(pat, val_str)
                        for run in element.runs:
                            run.text = ""
                        if element.runs:
                            element.runs[0].text = new_text
                        else:
                            element.add_run(new_text)

    def render(self, template_name: str, context: Dict[str, Any], user_id: str = "system") -> Dict[str, str]:
        from docx import Document
        
        template_path = os.path.join(self.template_dir, template_name)
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template {template_name} not found")

        docx_output_path = self._get_output_path(template_name.replace(".docx", ""), "docx")
        
        # Load Document
        doc = Document(template_path)
        
        # 1. Replace in Paragraphs (Body)
        for p in doc.paragraphs:
            self.replace_text_in_element(p, context)

        # 2. Replace in Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self.replace_text_in_element(p, context)
        
        # Save DOCX
        doc.save(docx_output_path)
        
        # Convert PDF
        pdf_output_path = None
        try:
            cmd_check = ["soffice", "--version"]
            subprocess.run(cmd_check, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", self.today_dir, docx_output_path]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            # LibreOffice output name logic
            expected_pdf = docx_output_path.replace(".docx", ".pdf")
            if os.path.exists(expected_pdf):
                pdf_output_path = expected_pdf
        except:
            logger.warning("PDF Conversion failed or LibreOffice not present")

        result = {
            "docx": docx_output_path,
            "pdf": pdf_output_path,
            "timestamp": datetime.datetime.now().isoformat(),
            "user": user_id
        }
        
        # Audit Log
        with open("form_engine/logs/audit.jsonl", "a", encoding='utf-8') as f:
            f.write(json.dumps(result, ensure_ascii=False) + "\n")

        return result
