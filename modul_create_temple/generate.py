#!/usr/bin/env python3
"""
Unified Form Generator - Replaces 18 individual generate_*b_only.py files.

Usage:
    python generate.py 1b                    # Generate form 1b
    python generate.py 3b --rejected         # Generate form 3b with rejected status
    python generate.py --all                 # Generate all forms
    python generate.py --list                # List available forms

Author: Refactored by Barry (Quick Flow Solo Dev)
"""
import sys
import os
import argparse
import datetime
import re
import subprocess
from pathlib import Path

# Add path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'form_engine/src'))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from form_engine.src.core.engine import (
    FormEngine,
    set_left_align_for_lists,
    CHECKBOX_CHECKED,
    CHECKBOX_UNCHECKED,
)
from sample_data import get_sample_data, TABLE_DATA_4B


# === CONFIGURATION ===
TEMPLATES_DIR = Path(__file__).parent / "form_engine/templates"
OUTPUT_DIR = Path(__file__).parent / "form_engine/output" / datetime.datetime.now().strftime("%Y-%m-%d")

# Forms that use simple engine.render()
SIMPLE_FORMS = ["1b", "2b", "7b", "8b", "9b", "11b", "12b", "13b", "18b", "pl1", "pl2", "pl3"]

# Forms that need cleanup logic (remove paragraphs based on approval status)
CLEANUP_FORMS = ["3b", "6b"]

# Forms with complex handling
COMPLEX_FORMS = ["4b", "5b", "10b", "14b", "15b", "16b", "17b"]

ALL_FORMS = SIMPLE_FORMS + CLEANUP_FORMS + COMPLEX_FORMS


def ensure_output_dir():
    """Create output directory if not exists."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def convert_to_pdf(docx_path: str) -> bool:
    """Convert DOCX to PDF using LibreOffice."""
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUTPUT_DIR), docx_path]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return True
    except Exception:
        return False


def verify_filled_variables(docx_path: str) -> list:
    """Check for unfilled variables in the generated document."""
    doc = Document(docx_path)
    text = ""
    for p in doc.paragraphs:
        text += p.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text += p.text + "\n"
    return list(set(re.findall(r'\{\{[^}]+\}\}', text)))


def replace_all_variables(doc: Document, engine: FormEngine, context: dict):
    """Replace variables in all document elements."""
    # Body paragraphs
    for p in doc.paragraphs:
        engine.replace_text_in_element(p, context)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    engine.replace_text_in_element(p, context)

    # Headers and Footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs:
            engine.replace_text_in_element(p, context)


def cleanup_approval_paragraphs(doc: Document, is_approved: bool):
    """Remove paragraphs based on approval status (for 3b, 6b forms)."""
    KEY_POSITIVE = "Đề nghị Nhà trường cho phép thực hiện"
    KEY_NEGATIVE = "Đề nghị Nhà trường không cho phép thực hiện"
    KEY_OR = "Hoặc"

    paragraphs_to_delete = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == KEY_OR:
            paragraphs_to_delete.append(p)
            continue
        if is_approved:
            if KEY_NEGATIVE in text or "không phù hợp sau" in text or "nội dung không phù hợp" in text:
                paragraphs_to_delete.append(p)
        else:
            if KEY_POSITIVE in text or "chỉnh sửa, bổ sung" in text or "nội dung chỉnh sửa" in text or "nội dung bổ sung" in text:
                paragraphs_to_delete.append(p)

    for p in paragraphs_to_delete:
        if p._element.getparent() is not None:
            p._element.getparent().remove(p._element)


def clean_empty_lines(doc: Document):
    """Remove empty lines and leftover template tags."""
    for i in range(len(doc.paragraphs) - 1, 0, -1):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if not text or text in [".", ",", "_"]:
            try:
                p._element.getparent().remove(p._element)
            except:
                pass
        if "}}" in p.text:
            p.text = p.text.replace("}}", "")


def set_cell_format(cell, text, bold=False, align=None):
    """Format a table cell."""
    cell.text = text
    for paragraph in cell.paragraphs:
        if align:
            paragraph.alignment = align
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = bold


def set_cell_border(cell):
    """Add borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        element = OxmlElement(f'w:{border}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        tcBorders.append(element)
    tcPr.append(tcBorders)


# === FORM-SPECIFIC HANDLERS ===

def generate_simple(form_id: str, is_approved: bool) -> str:
    """Generate simple forms using engine.render()."""
    engine = FormEngine()
    context = get_sample_data(form_id, is_approved)
    template_name = f"{form_id.upper() if form_id.startswith('pl') else form_id}.docx"

    result = engine.render(template_name, context, user_id="unified_generator")
    return result['docx']


def generate_with_cleanup(form_id: str, is_approved: bool) -> str:
    """Generate forms that need cleanup logic (3b, 6b)."""
    context = get_sample_data(form_id, is_approved)
    template_path = TEMPLATES_DIR / f"{form_id}.docx"
    output_path = OUTPUT_DIR / f"{form_id}_{'approved' if is_approved else 'rejected'}.docx"

    doc = Document(str(template_path))
    engine = FormEngine()

    # Replace variables
    replace_all_variables(doc, engine, context)

    # Cleanup based on approval
    cleanup_approval_paragraphs(doc, is_approved)
    clean_empty_lines(doc)

    doc.save(str(output_path))
    convert_to_pdf(str(output_path))
    return str(output_path)


def generate_4b(is_approved: bool) -> str:
    """Generate form 4b with dynamic table."""
    context = get_sample_data("4b", is_approved)
    context["so_de_tai"] = str(len(TABLE_DATA_4B))
    template_path = TEMPLATES_DIR / "4b.docx"
    output_path = OUTPUT_DIR / "4b_generated.docx"

    doc = Document(str(template_path))

    # Fill variables in all elements
    for p in doc.paragraphs:
        for k, v in context.items():
            if f"{{{{{k}}}}}" in p.text:
                p.text = p.text.replace(f"{{{{{k}}}}}", str(v))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in context.items():
                        if f"{{{{{k}}}}}" in p.text:
                            p.text = p.text.replace(f"{{{{{k}}}}}", str(v))

    # Fill data table
    if len(doc.tables) > 1:
        target_table = doc.tables[1]
        for item in TABLE_DATA_4B:
            row = target_table.add_row()
            cells = row.cells
            set_cell_format(cells[0], str(item['stt']), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_border(cells[0])
            set_cell_format(cells[1], item['ten'])
            set_cell_border(cells[1])
            set_cell_format(cells[2], item['cn'])
            set_cell_border(cells[2])
            if len(cells) > 3:
                set_cell_format(cells[3], item['mt'])
                set_cell_border(cells[3])
            if len(cells) > 4:
                set_cell_format(cells[4], item['tm'])
                set_cell_border(cells[4])
            if len(cells) > 5:
                set_cell_format(cells[5], "- Nội dung 1\n- Nội dung 2")
                set_cell_border(cells[5])
            if len(cells) > 6:
                set_cell_format(cells[6], item['kq'])
                set_cell_border(cells[6])
            if len(cells) > 7:
                set_cell_format(cells[7], item['ud'])
                set_cell_border(cells[7])
            if len(cells) > 8:
                set_cell_format(cells[8], item['kp'], align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_border(cells[8])

        # Remove template row
        if len(target_table.rows) > 1:
            row_1_text = target_table.rows[1].cells[0].text
            if "{{" in row_1_text:
                tbl = target_table._tbl
                tr = target_table.rows[1]._tr
                tbl.remove(tr)

    doc.save(str(output_path))
    convert_to_pdf(str(output_path))
    return str(output_path)


def generate_10b(is_approved: bool) -> str:
    """Generate form 10b with complex paragraph handling."""
    context = get_sample_data("10b", is_approved)
    template_path = TEMPLATES_DIR / "10b.docx"
    output_path = OUTPUT_DIR / f"10b_{'approved' if is_approved else 'rejected'}.docx"

    doc = Document(str(template_path))
    engine = FormEngine()

    # Find section 8 positions
    section8_idx = -1
    phien_hop_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "Kết luận của Hội đồng":
            section8_idx = i
        if "Phiên họp kết thúc" in text:
            phien_hop_idx = i
            break

    # Process section 8
    if section8_idx >= 0 and phien_hop_idx > section8_idx:
        doc.paragraphs[section8_idx].text = "8. Kết luận của Hội đồng"
        for run in doc.paragraphs[section8_idx].runs:
            run.bold = True
            run.font.name = 'Times New Roman'

        # Remove template paragraphs
        for idx in range(phien_hop_idx - 1, section8_idx, -1):
            p_element = doc.paragraphs[idx]._element
            p_element.getparent().remove(p_element)

        # Find "Phiên họp kết thúc" again
        for i, p in enumerate(doc.paragraphs):
            if "Phiên họp kết thúc" in p.text:
                phien_hop_idx = i
                break

        # Insert conclusion content
        if phien_hop_idx >= 0:
            if is_approved:
                texts = [
                    (context.get('noi_dung_bo_sung', ''), True),
                    ("- Những nội dung bổ sung:", False),
                    (context.get('noi_dung_da_chinh_sua', ''), True),
                    ("- Những nội dung chỉnh sửa:", False),
                    (f"Đề nghị Nhà trường cho phép thực hiện đề tài \"{context['ten_de_tai']}\" sau khi chỉnh sửa:", False),
                ]
            else:
                texts = [
                    (context.get('noi_dung_khong_phu_hop', ''), True),
                    (f"Đề nghị Nhà trường không cho phép thực hiện đề tài \"{context['ten_de_tai']}\":", False),
                ]

            for text, is_multiline in texts:
                if is_multiline and '\n' in str(text):
                    for line in reversed(str(text).split('\n')):
                        if line.strip():
                            new_para = doc.paragraphs[phien_hop_idx].insert_paragraph_before(line)
                            new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif str(text).strip():
                    new_para = doc.paragraphs[phien_hop_idx].insert_paragraph_before(str(text))
                    new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Replace remaining variables
    replace_all_variables(doc, engine, context)
    set_left_align_for_lists(doc)

    # Clean leftover tags
    for p in doc.paragraphs:
        if '{{' in p.text:
            p.text = re.sub(r'\{\{[^\}]*\}\}', '', p.text)

    doc.save(str(output_path))
    convert_to_pdf(str(output_path))
    return str(output_path)


def generate_14b(is_approved: bool) -> str:
    """Generate form 14b."""
    context = get_sample_data("14b", is_approved)
    template_path = TEMPLATES_DIR / "14b.docx"
    output_path = OUTPUT_DIR / f"14b_{'approved' if is_approved else 'rejected'}.docx"

    doc = Document(str(template_path))
    engine = FormEngine()

    replace_all_variables(doc, engine, context)
    set_left_align_for_lists(doc)

    doc.save(str(output_path))
    convert_to_pdf(str(output_path))
    return str(output_path)


def generate_complex(form_id: str, is_approved: bool) -> str:
    """Generate complex forms (5b, 15b, 16b, 17b) with standard processing."""
    context = get_sample_data(form_id, is_approved)
    template_path = TEMPLATES_DIR / f"{form_id}.docx"
    output_path = OUTPUT_DIR / f"{form_id}_generated.docx"

    doc = Document(str(template_path))
    engine = FormEngine()

    replace_all_variables(doc, engine, context)
    set_left_align_for_lists(doc)

    doc.save(str(output_path))
    convert_to_pdf(str(output_path))
    return str(output_path)


# === MAIN DISPATCHER ===

def generate_form(form_id: str, is_approved: bool = True) -> str:
    """Main dispatcher for form generation."""
    form_id = form_id.lower()
    ensure_output_dir()

    print(f"--- GENERATING FORM {form_id.upper()} ({'APPROVED' if is_approved else 'REJECTED'}) ---")

    try:
        if form_id in SIMPLE_FORMS:
            output = generate_simple(form_id, is_approved)
        elif form_id in CLEANUP_FORMS:
            output = generate_with_cleanup(form_id, is_approved)
        elif form_id == "4b":
            output = generate_4b(is_approved)
        elif form_id == "10b":
            output = generate_10b(is_approved)
        elif form_id == "14b":
            output = generate_14b(is_approved)
        elif form_id in COMPLEX_FORMS:
            output = generate_complex(form_id, is_approved)
        else:
            print(f"❌ Unknown form: {form_id}")
            return ""

        # Verify
        unfilled = verify_filled_variables(output)
        if unfilled:
            print(f"⚠️  {len(unfilled)} unfilled variables: {', '.join(unfilled[:5])}")
        else:
            print("✅ All variables filled!")

        print(f"📄 Output: {output}")
        return output

    except Exception as e:
        print(f"❌ Error generating {form_id}: {e}")
        import traceback
        traceback.print_exc()
        return ""


def main():
    parser = argparse.ArgumentParser(description="Unified Form Generator")
    parser.add_argument("form_id", nargs="?", help="Form ID (e.g., 1b, 2b, 3b, pl1)")
    parser.add_argument("--rejected", action="store_true", help="Generate with rejected status")
    parser.add_argument("--all", action="store_true", help="Generate all forms")
    parser.add_argument("--list", action="store_true", help="List available forms")

    args = parser.parse_args()

    if args.list:
        print("Available forms:")
        print(f"  Simple:  {', '.join(SIMPLE_FORMS)}")
        print(f"  Cleanup: {', '.join(CLEANUP_FORMS)}")
        print(f"  Complex: {', '.join(COMPLEX_FORMS)}")
        return

    if args.all:
        print("=== GENERATING ALL FORMS ===\n")
        for form_id in ALL_FORMS:
            generate_form(form_id, is_approved=True)
            print()
        return

    if not args.form_id:
        parser.print_help()
        return

    is_approved = not args.rejected
    generate_form(args.form_id, is_approved)


if __name__ == "__main__":
    main()
