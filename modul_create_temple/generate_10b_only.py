import sys
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import (
    FormEngine, set_left_align_for_lists,
    CHECKBOX_CHECKED, CHECKBOX_UNCHECKED
)

def main():
    print("--- GENERATING FORM 10B (BIÊN BẢN HỌP HỘI ĐỒNG) ---")

    # 1. Prepare Data
    context = {
        # Thông tin đề tài
        "qd_so": "123/QĐ-ĐHSPKT",
        "ma_de_tai": "NCKH-2024-01",
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "chu_nhiem_de_tai": "TS. Nguyễn Văn AI",

        # Thông tin cuộc họp
        "don_vi_chu_tri": "Khoa Công nghệ thông tin",
        "dia_diem": "Phòng họp Khoa CNTT",
        "thoi_gian_hop": "14:00 ngày 30/01/2024",
        "thoi_gian_ket_thuc": "16:30 ngày 30/01/2024",

        # Thành phần tham dự
        "ten_chu_tich": "PGS.TS. Trần Văn B",
        "ten_thu_ky": "ThS. Lê Thị C",
        "thanh_vien": (
            "- TS. Phạm Văn C - Khoa Điện\n"
            "- TS. Hoàng Văn D - Khoa Cơ khí\n"
            "- ThS. Ngô Thị E - Trung tâm Số"
        ),

        # Số liệu thành viên
        "tong_phieu_phat_ra": "5",
        "tong_phieu_thu_vao": "5",
        "so_phieu_dat": "5",
        "so_phieu_khong_dat": "0",

        # Điểm danh
        "co_mat_tren_tong": "5/5",
        "vang_mat": "0",

        # Ngày họp
        "ngay_ht": "30",

        # Nội dung đánh giá
        "cac_chi_tieu_chu_yeu_cac_yeu_cau_khoa_hoc_cua_ket_qua": (
            "- Đề tài đáp ứng các chỉ tiêu khoa học cơ bản.\n"
            "- Phương pháp nghiên cứu phù hợp, kết quả có giá trị thực tiễn.\n"
            "- Số lượng bài báo khoa học: 02 bài đăng trên tạp chí quốc tế."
        ),

        "phuong_phap_nghien_cuu": (
            "- Khảo sát thực trạng tại 3 đơn vị đại học.\n"
            "- Thiết kế và xây dựng hệ thống Blockchain.\n"
            "- Kiểm thử và đánh giá hiệu quả."
        ),

        "so_luong_chung_loai_khoi_luong_san_pham": (
            "- 01 hệ thống quản lý văn bằng trên nền Blockchain.\n"
            "- 02 bài báo khoa học (01 quốc tế, 01 trong nước)."
        ),

        "noi_dung_da_chinh_sua": (
            "- Đã chỉnh sửa lại tên đề tài theo kết luận của Hội đồng.\n"
            "- Đã làm rõ hơn phương pháp nghiên cứu định lượng."
        ),

        "noi_dung_bo_sung": (
            "- Đã bổ sung 01 thành viên nghiên cứu là sinh viên.\n"
            "- Đã thêm kế hoạch khảo sát thực tế."
        ),

        "noi_dung_khong_phu_hop": "",  # Để trống khi không có nội dung không phù hợp

        "nhan_xet_ve_muc_do_hoan_thanh": (
            "Đề tài đã hoàn thành đầy đủ các nội dung theo đề xuất ban đầu. "
            "Kết quả nghiên cứu có tính ứng dụng cao, có thể triển khai thực tế."
        ),

        # Checkbox - QUY TẮC CHUNG
        "box_dat": CHECKBOX_CHECKED,
        "box_khong_dat": CHECKBOX_UNCHECKED,
    }

    # 2. Render
    tmpl_path = "form_engine/templates/10b.docx"
    out_docx = "form_engine/output/2026-01-16/10b_generated.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()

    # PHƯƠNG ANH 10B CẦN XỬ LÝ ĐẶC BIỆT:
    # Template có block điều kiện phức tạp. Ta xử lý TRƯỚC KHI replace.

    # Tìm positions TRƯỚC KHI sửa gì
    # Template structure:
    #   Para 14: "Kết luận của Hội đồng:" (có dấu :) - Section 7
    #   Para 31: "Kết luận của Hội đồng" (không có :) - Section 8
    #   Para 32-41: Block template phức tạp (cần xóa)
    #   Para 42: "Phiên họp kết thúc..."
    section8_idx = -1
    phien_hop_idx = -1

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # Tìm "Kết luận của Hội đồng" KHÔNG có dấu : (đó là Section 8)
        if text == "Kết luận của Hội đồng":
            section8_idx = i
        if "Phiên họp kết thúc" in text:
            phien_hop_idx = i
            break  # Tìm thấy là dừng

    # Xử lý mục 8 và xóa block template phức tạp
    if section8_idx >= 0 and phien_hop_idx > section8_idx:
        # 1. Sửa tên thành "8. Kết luận của Hội đồng"
        doc.paragraphs[section8_idx].text = "8. Kết luận của Hội đồng"
        for run in doc.paragraphs[section8_idx].runs:
            run.bold = True  # Tiêu đề mục thì in đậm
            run.font.name = 'Times New Roman'

        # 2. Xóa paragraphs từ sau mục 8 đến trước "Phiên họp kết thúc"
        # Lưu ý: Sau khi xóa, các index sẽ thay đổi, nên phải xóa từ dưới lên
        for idx in range(phien_hop_idx - 1, section8_idx, -1):
            p_element = doc.paragraphs[idx]._element
            p_element.getparent().remove(p_element)

        # Tìm lại "Phiên họp kết thúc" sau khi xóa
        phien_hop_idx = -1
        for i, p in enumerate(doc.paragraphs):
            if "Phiên họp kết thúc" in p.text:
                phien_hop_idx = i
                break

        # Insert nội dung mới TRƯỚC "Phiên họp kết thúc"
        if phien_hop_idx >= 0:
            is_pass = context.get("box_dat") == CHECKBOX_CHECKED

            if is_pass:
                # Insert theo thứ tự NGƯỢC để đúng hiển thị
                # Các multi-line string cần được tách thành paragraphs riêng
                texts_and_multiline = [
                    (context['noi_dung_bo_sung'], True),  # multi-line
                    ("- Những nội dung bổ sung:", False),
                    (context['noi_dung_da_chinh_sua'], True),  # multi-line
                    ("- Những nội dung chỉnh sửa:", False),
                    (f"Đề nghị Nhà trường cho phép thực hiện đề tài \"{context['ten_de_tai']}\" sau khi Nhóm tác giả đã chỉnh sửa, bổ sung, hoàn thiện hồ sơ theo góp ý của Hội đồng như sau:", False),
                ]
            else:
                texts_and_multiline = [
                    (context['noi_dung_khong_phu_hop'], True),
                    (f"Đề nghị Nhà trường không cho phép thực hiện đề tài \"{context['ten_de_tai']}\" vì những nội dung không phù hợp sau:", False),
                ]

            # Insert trước "Phiên họp kết thúc"
            for text, is_multiline in texts_and_multiline:
                if is_multiline and '\n' in text:
                    # Tách theo \n và insert từng dòng, theo thứ tự NGƯỢC
                    lines = text.split('\n')
                    for line in reversed(lines):
                        if line.strip():  # Chỉ insert non-empty lines
                            new_para = doc.paragraphs[phien_hop_idx].insert_paragraph_before(line)
                            for run in new_para.runs:
                                run.bold = False
                                run.font.name = 'Times New Roman'
                            new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    if text.strip():  # Chỉ insert non-empty
                        new_para = doc.paragraphs[phien_hop_idx].insert_paragraph_before(text)
                        for run in new_para.runs:
                            run.bold = False
                            run.font.name = 'Times New Roman'
                        new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # BƯỚC 2: Replace các biến còn lại
    for p in doc.paragraphs:
        engine.replace_text_in_element(p, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    engine.replace_text_in_element(p, context)

    for section in doc.sections:
        for p in section.header.paragraphs:
            engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs:
            engine.replace_text_in_element(p, context)

    # 3. Căn trái các danh sách gạch đầu dòng (QUY TẮC CHUNG)
    set_left_align_for_lists(doc)

    # 4. Xóa các tag {{...}} còn sót
    for p in doc.paragraphs:
        if '{{' in p.text:
            new_text = re.sub(r'\{\{[^\}]*\}\}', '', p.text)
            if new_text != p.text:
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '{{' in p.text:
                        new_text = re.sub(r'\{\{[^\}]*\}\}', '', p.text)
                        if new_text != p.text:
                            for run in p.runs:
                                run.text = ""
                            if p.runs:
                                p.runs[0].text = new_text

    # 5. Tìm lại và sửa mục 8 sau khi replace (để đảm bảo đúng)
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "Kết luận của Hội đồng" and i > 20:  # Dòng thứ 2
            p.text = "8. Kết luận của Hội đồng"
            for run in p.runs:
                run.bold = False
                run.font.name = 'Times New Roman'
            break

    # Save
    doc.save(out_docx)

    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 10b: {out_docx.replace('.docx', '.pdf')}")
    except:
        print(f"⚠️ PDF conversion failed")

    # Verify unfilled variables
    doc_check = Document(out_docx)
    text = ""
    for p in doc_check.paragraphs:
        text += p.text + "\n"
    for table in doc_check.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text += p.text + "\n"

    unfilled = re.findall(r'\{\{[^}]+\}\}', text)
    if unfilled:
        print(f"⚠️ Còn {len(set(unfilled))} biến chưa điền:")
        for u in sorted(set(unfilled)):
            print(f"     {u}")
    else:
        print("✅ Tất cả biến đã điền đầy đủ!")

if __name__ == "__main__":
    main()
