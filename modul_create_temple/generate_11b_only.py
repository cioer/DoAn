import sys
import os
from docx import Document

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine, set_left_align_for_lists

def main():
    print("--- GENERATING FORM 11B (BÁO CÁO HOÀN THIỆN HỒ SƠ) ---")

    # 1. Prepare Data
    context = {
        # Thông tin chung
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "nguoi_de_xuat": "TS. Nguyễn Văn AI",
        "ten_nguoi_de_xuat": "TS. Nguyễn Văn AI",
        "thoi_gian_hop": "14:00 ngày 30/01/2024",

        # Nội dung đã chỉnh sửa
        "noi_dung_da_sua": (
            "- Đã chỉnh sửa lại tên đề tài theo kết luận của Hội đồng.\n"
            "- Đã làm rõ hơn phương pháp nghiên cứu định lượng.\n"
            "- Đã bổ sung thêm số liệu thống kê."
        ),

        # Nội dung bổ sung
        "noi_dung_bo_sung": (
            "- Đã bổ sung 01 thành viên nghiên cứu là sinh viên.\n"
            "- Đã thêm kế hoạch khảo sát thực tế.\n"
            "- Đã hoàn thiện báo cáo tổng hợp theo yêu cầu."
        ),

        # Người ký
        "ten_chu_tich_hoi_dong": "PGS.TS. Trần Văn B",
        "ten_thu_ky": "ThS. Lê Thị C",

        # Ngày tháng
        "ngay": "20",
        "thang": "01",
        "nam": "2024",
    }

    # 2. Render
    tmpl_path = "form_engine/templates/11b.docx"
    out_docx = "form_engine/output/2026-01-16/11b_generated.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()

    # Replace Variables in ALL places (QUY TẮC CHUNG)
    for p in doc.paragraphs:
        engine.replace_text_in_element(p, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    engine.replace_text_in_element(p, context)

    for section in doc.sections:
        for p in section.header.paragraphs:
            engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs:
            engine.replace_text_in_element(p, context)

    # 3. Căn trái các danh sách gạch đầu dòng (QUY TẮC CHUNG)
    set_left_align_for_lists(doc)

    # Save
    doc.save(out_docx)

    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 11b: {out_docx.replace('.docx', '.pdf')}")
    except:
        print(f"⚠️ PDF conversion failed")

    # Verify unfilled variables
    doc_check = Document(out_docx)
    import re
    text = ""
    for p in doc_check.paragraphs:
        text += p.text + "\n"
    for table in doc_check.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text += p.text + "\n"

    unfilled = re.findall(r'\{\{[^}]+\}\}', text)
    if unfilled:
        print(f"⚠️ Còn {len(set(unfilled))} biến chưa điền:")
        for u in sorted(set(unfilled)):
            print(f"     {u}")
    else:
        print("✅ Tất cả biến đã điền đầy đủ!")

if __name__ == "__main__":
    main()
