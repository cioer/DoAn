import sys
import os
from docx import Document

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine, set_left_align_for_lists

def main():
    print("--- GENERATING FORM 12B (NHẬN XÉT PHẢN BIỆN) ---")

    # 1. Prepare Data
    context = {
        # Quyết định
        "qd_so": "123/QĐ-ĐHSPKT",
        "ngay": "15",
        "thang": "01",
        "nam": "2024",

        # Thông tin đề tài
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "ma_so_de_tai": "NCKH-2024-01",

        # Chủ nhiệm
        "ho_ten_chu_nhiem": "TS. Nguyễn Văn AI",
        "don_vi": "Khoa Công nghệ thông tin",

        # Thành viên
        "nhung_thanh_vien": (
            "- ThS. Trần Thị B - Khoa CNTT\n"
            "- TS. Lê Văn C - Trung tâm Số"
        ),

        # Người phản biện
        "ten_nguoi_phan_bien": "PGS.TS. Hoàng Văn D",
        "don_vi_nguoi_phan_dien": "Khoa Điện",  # Template có typo "dien"
        "don_vi_nguoi_phan_bien": "Khoa Điện",  # Fallback đúng chính tả

        # Kiến nghị và kết luận
        "kien_nghi_ket_luan": (
            "Đề tài có tính cấp thiết cao, phương pháp nghiên cứu khoa học. "
            "Kết quả nghiên cứu có giá trị thực tiễn, có thể ứng dụng trong quản lý văn bằng. "
            "Tài liệu tham khảo đầy đủ, cập nhật. "
            "Đề nghị Hội đồng đánh giá, nghiệm thu đề tài."
        ),

        # Người nhận xét
        "ten_nguoi_nhan_xet": "PGS.TS. Hoàng Văn D",
    }

    # 2. Render
    tmpl_path = "form_engine/templates/12b.docx"
    out_docx = "form_engine/output/2026-01-16/12b_generated.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()

    # Replace Variables in ALL places (QUY TẮC CHUNG)
    for p in doc.paragraphs:
        engine.replace_text_in_element(p, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    engine.replace_text_in_element(p, context)

    for section in doc.sections:
        for p in section.header.paragraphs:
            engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs:
            engine.replace_text_in_element(p, context)

    # 3. Căn trái các danh sách gạch đầu dòng (QUY TẮC CHUNG)
    set_left_align_for_lists(doc)

    # Save
    doc.save(out_docx)

    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 12b: {out_docx.replace('.docx', '.pdf')}")
    except:
        print(f"⚠️ PDF conversion failed")

    # Verify unfilled variables
    doc_check = Document(out_docx)
    import re
    text = ""
    for p in doc_check.paragraphs:
        text += p.text + "\n"
    for table in doc_check.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text += p.text + "\n"

    unfilled = re.findall(r'\{\{[^}]+\}\}', text)
    if unfilled:
        print(f"⚠️ Còn {len(set(unfilled))} biến chưa điền:")
        for u in sorted(set(unfilled)):
            print(f"     {u}")
    else:
        print("✅ Tất cả biến đã điền đầy đủ!")

if __name__ == "__main__":
    main()
