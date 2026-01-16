import sys
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import (
    FormEngine, set_left_align_for_lists,
    fill_cell_text, CHECKBOX_CHECKED, CHECKBOX_UNCHECKED
)

def main():
    print("--- GENERATING FORM 13B (GIẤY ĐỀ NGHỊ THÀNH LẬP HĐ) ---")

    # 1. Prepare Data
    context = {
        # Ngày tháng
        "ngay": "20",
        "thang": "01",
        "nam": "2024",

        # Khoa
        "ten_khoa": "Công nghệ thông tin",

        # Thông tin đề tài
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "ma_de_tai": "NCKH-2024-01",

        # Chủ nhiệm đề tài
        "chu_nhiem": "TS. Nguyễn Văn AI",

        # Hội đồng - Chủ tịch và Thư ký
        "ten_chu_tich": "PGS.TS. Trần Văn B",
        "don_vi_cong_tac_chu_tich": "Khoa Công nghệ thông tin",

        "ten_thu_ky": "ThS. Lê Thị C",
        "don_vi_cong_tac_thu_ky": "Khoa Công nghệ thông tin",

        # Danh sách ủy viên (STT bắt đầu từ 3)
        "hoi_dong_uy_vien": [
            {"stt": 3, "ho_ten": "TS. Phạm Văn D", "don_vi": "Khoa Điện"},
            {"stt": 4, "ho_ten": "TS. Hoàng Văn E", "don_vi": "Khoa Cơ khí"},
            {"stt": 5, "ho_ten": "ThS. Ngô Thị F", "don_vi": "Trung tâm Số"},
        ],
    }

    # 2. Render
    tmpl_path = "form_engine/templates/13b.docx"
    out_docx = "form_engine/output/2026-01-16/13b_generated.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()

    # XỬ LÝ BẢNG HỘI ĐỒNG TRƯỚC (Table 1)
    # Table 1 có cấu trúc:
    #   Row 0: Header
    #   Row 1: Chủ tịch (STT 1)
    #   Row 2: Thư ký (STT 2)
    #   Row 3: Template row cho ủy viên (cần duplicate)
    hoi_dong_table = doc.tables[1]

    # Fill Row 1: Chủ tịch
    fill_cell_text(hoi_dong_table.rows[1].cells[0], "1", WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell_text(hoi_dong_table.rows[1].cells[1], context["ten_chu_tich"], WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell_text(hoi_dong_table.rows[1].cells[2], context["don_vi_cong_tac_chu_tich"], WD_ALIGN_PARAGRAPH.LEFT)

    # Fill Row 2: Thư ký
    fill_cell_text(hoi_dong_table.rows[2].cells[0], "2", WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell_text(hoi_dong_table.rows[2].cells[1], context["ten_thu_ky"], WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell_text(hoi_dong_table.rows[2].cells[2], context["don_vi_cong_tac_thu_ky"], WD_ALIGN_PARAGRAPH.LEFT)

    # Xóa template row (Row 3) và add rows cho ủy viên
    # Lưu template row structure để copy format
    template_row = hoi_dong_table.rows[3]

    # Xóa template row
    template_row._element.getparent().remove(template_row._element)

    # Thêm rows cho ủy viên (STT từ 3 trở đi)
    for uv in context["hoi_dong_uy_vien"]:
        new_row = hoi_dong_table.add_row()
        fill_cell_text(new_row.cells[0], str(uv["stt"]), WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell_text(new_row.cells[1], uv["ho_ten"], WD_ALIGN_PARAGRAPH.LEFT)
        fill_cell_text(new_row.cells[2], uv["don_vi"], WD_ALIGN_PARAGRAPH.LEFT)
        fill_cell_text(new_row.cells[3], "Ủy viên", WD_ALIGN_PARAGRAPH.LEFT)

    # Replace Variables in ALL places (QUY TẮC CHUNG)
    for p in doc.paragraphs:
        engine.replace_text_in_element(p, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    engine.replace_text_in_element(p, context)

    for section in doc.sections:
        for p in section.header.paragraphs:
            engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs:
            engine.replace_text_in_element(p, context)

    # 3. Căn trái các danh sách gạch đầu dòng (QUY TẮC CHUNG)
    set_left_align_for_lists(doc)

    # Save
    doc.save(out_docx)

    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 13b: {out_docx.replace('.docx', '.pdf')}")
    except:
        print(f"⚠️ PDF conversion failed")

    # Verify unfilled variables
    doc_check = Document(out_docx)
    import re
    text = ""
    for p in doc_check.paragraphs:
        text += p.text + "\n"
    for table in doc_check.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text += p.text + "\n"

    unfilled = re.findall(r'\{\{[^}]+\}\}', text)
    if unfilled:
        print(f"⚠️ Còn {len(set(unfilled))} biến chưa điền:")
        for u in sorted(set(unfilled)):
            print(f"     {u}")
    else:
        print("✅ Tất cả biến đã điền đầy đủ!")

if __name__ == "__main__":
    main()
