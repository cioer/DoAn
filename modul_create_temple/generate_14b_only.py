import sys
import os
from docx import Document

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine, set_left_align_for_lists, CHECKBOX_CHECKED, CHECKBOX_UNCHECKED

def main():
    print("--- GENERATING FORM 14B (PHIẾU ĐÁNH GIÁ NGHIỆM THU CẤP TRƯỜNG) ---")

    # 1. Prepare Data
    is_pass = True  # Đạt = True, Không đạt = False

    context = {
        # Thông tin người đánh giá
        "ho_ten": "TS. Phạm Văn D",
        "don_vi_cong_tac": "Khoa Điện",

        # Thông tin đề tài
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "ma_so_de_tai": "NCKH-2024-01",
        "ten_chu_nhiem": "TS. Nguyễn Văn AI",
        "don_vi_chu_tri": "Khoa Công nghệ thông tin",

        # Thời gian, địa điểm
        "thoi_gian_hop": "14:00 ngày 30/01/2024",
        "dia_diem": "Phòng họp Trường",

        # Checkbox cho bảng đánh giá (QUY TẮC CHUNG: [x]/[ ])
        "dat": CHECKBOX_CHECKED if is_pass else CHECKBOX_UNCHECKED,
        "khong_dat": CHECKBOX_UNCHECKED if is_pass else CHECKBOX_CHECKED,

        # Checkbox kết luận chung
        "box_dat": CHECKBOX_CHECKED if is_pass else CHECKBOX_UNCHECKED,
        "box_khong_dat": CHECKBOX_UNCHECKED if is_pass else CHECKBOX_CHECKED,

        # Ý kiến khác
        "y_kien_khac": "Đề nghị chủ nhiệm đề tài tiếp tục phát triển và mở rộng hệ thống.",

        # Người đánh giá
        "ten_nguoi_danh_gia": "TS. Phạm Văn D",
    }

    # 2. Render
    tmpl_path = "form_engine/templates/14b.docx"
    out_docx = "form_engine/output/2026-01-16/14b_generated.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()

    # Replace Variables in ALL places (QUY TẮC CHUNG)
    for p in doc.paragraphs:
        engine.replace_text_in_element(p, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    engine.replace_text_in_element(p, context)

    for section in doc.sections:
        for p in section.header.paragraphs:
            engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs:
            engine.replace_text_in_element(p, context)

    # 3. Căn trái các danh sách gạch đầu dòng (QUY TẮC CHUNG)
    set_left_align_for_lists(doc)

    # Save
    doc.save(out_docx)

    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 14b: {out_docx.replace('.docx', '.pdf')}")
    except:
        print(f"⚠️ PDF conversion failed")

    # Verify unfilled variables
    doc_check = Document(out_docx)
    import re
    text = ""
    for p in doc_check.paragraphs:
        text += p.text + "\n"
    for table in doc_check.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text += p.text + "\n"

    unfilled = re.findall(r'\{\{[^}]+\}\}', text)
    if unfilled:
        print(f"⚠️ Còn {len(set(unfilled))} biến chưa điền:")
        for u in sorted(set(unfilled)):
            print(f"     {u}")
    else:
        print("✅ Tất cả biến đã điền đầy đủ!")

if __name__ == "__main__":
    main()
