import sys
import os
from docx import Document

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine, set_left_align_for_lists

def main():
    print("--- GENERATING FORM 18B (ĐƠN XIN GIA HẠN) ---")

    # 1. Prepare Data
    context = {
        # Ngày tháng
        "ngay": "20",
        "thang": "01",
        "nam": "2024",

        # Khoa
        "ten_khoa": "Công nghệ thông tin",

        # Chủ nhiệm
        "ten_chu_nhiem": "TS. Nguyễn Văn AI",
        "ten_don_vi": "Khoa Công nghệ thông tin",

        # Thông tin đề tài
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "ma_so_de_tai": "NCKH-2024-01",

        # Quyết định
        "qd_so": "123/QĐ-ĐHSPKT",
        "ngay_ht": "15/01/2024",

        # Thời gian
        "thoi_gian_hoan_thanh": "12 tháng",
        "thoi_gian_xin_gia_han": "06 tháng",

        # Lý do
        "ly_do": "Do cần thêm thời gian để thu thập dữ liệu thực tế từ các đơn vị đối tác.",

        # Kết quả đã đạt được
        "da_dat_duoc_ket_qua_gi_so_voi_de_cuong": (
            "- Đã hoàn thành 80% nội dung nghiên cứu theo đề cương.\n"
            "- Đã thiết kế xong kiến trúc hệ thống Blockchain.\n"
            "- Đã triển khai thử nghiệm tại 01 đơn vị demo.\n"
            "- Đang hoàn thiện báo cáo tổng hợp."
        ),

        # Người ký (Table 1)
        "ten_truong_khoa": "TS. Phạm Văn X",
        "phong_KHCN": "TS. Trần Thị Y",
        "hoi_dong_truong": "PGS.TS. Nguyễn Văn Z",
    }

    # 2. Render
    tmpl_path = "form_engine/templates/18b.docx"
    out_docx = "form_engine/output/2026-01-16/18b_generated.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()

    # Replace Variables in ALL places (QUY TẮC CHUNG)
    for p in doc.paragraphs:
        engine.replace_text_in_element(p, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    engine.replace_text_in_element(p, context)

    for section in doc.sections:
        for p in section.header.paragraphs:
            engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs:
            engine.replace_text_in_element(p, context)

    # 3. Căn trái các danh sách gạch đầu dòng (QUY TẮC CHUNG)
    set_left_align_for_lists(doc)

    # Save
    doc.save(out_docx)

    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 18b: {out_docx.replace('.docx', '.pdf')}")
    except:
        print(f"⚠️ PDF conversion failed")

    # Verify unfilled variables
    doc_check = Document(out_docx)
    import re
    text = ""
    for p in doc_check.paragraphs:
        text += p.text + "\n"
    for table in doc_check.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text += p.text + "\n"

    unfilled = re.findall(r'\{\{[^}]+\}\}', text)
    if unfilled:
        print(f"⚠️ Còn {len(set(unfilled))} biến chưa điền:")
        for u in sorted(set(unfilled)):
            print(f"     {u}")
    else:
        print("✅ Tất cả biến đã điền đầy đủ!")

if __name__ == "__main__":
    main()
