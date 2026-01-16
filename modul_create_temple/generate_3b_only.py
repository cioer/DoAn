import sys
import os
from docx import Document

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine

def main():
    print("--- GENERATING FORM 3B (FINAL FORMAT) ---")
    
    # 1. DỮ LIỆU ĐẦU VÀO
    # Chủ nhiệm
    chu_nhiem = {"ten": "ThS. Phan Đức Thiện", "don_vi": "Trung tâm Thực hành"}
    
    # Thành viên
    hoi_dong = [
        {"ten": "ThS. Nguyễn Văn Trung", "don_vi": "Khoa CNTT"},
        {"ten": "ThS. Trần Thị Hồng Nhung", "don_vi": "Phòng KH-HTQT & ĐBCL"},
        {"ten": "ThS. Nguyễn Thị Thu Thủy", "don_vi": "Khoa CNTT"},
        {"ten": "ThS. Nguyễn Thế Vinh", "don_vi": "Khoa CNTT"}
    ]
    
    # 2. XỬ LÝ FORMATTING
    # Format Chủ nhiệm: • Chủ nhiệm: [Tên] [Tab] Đơn vị: [Đơn vị]
    # Lưu ý: Template có thể đã có chữ "Chủ nhiệm:" rồi. Cần check Template.
    # Nhưng theo yêu cầu của bạn, tôi sẽ format nguyên chuỗi này.
    # Nếu template là: "Chủ nhiệm: {{chu_nhiem}}" -> Tôi chỉ điền Tên.
    # Nếu template là: "{{thong_tin_nhom_tac_gia}}" -> Tôi điền cả cục.
    
    # Dựa vào lần scan trước, Template 3b có biến {{chu_nhiem}} và {{thanh_vien}}.
    # Giả sử Template chỉ có 2 biến này nằm trơ trọi.
    
    txt_chu_nhiem = f"• Chủ nhiệm: {chu_nhiem['ten']}\t\tĐơn vị: {chu_nhiem['don_vi']}"
    
    txt_thanh_vien = "• Thành viên:"
    for mem in hoi_dong:
        # Dùng \t để tab, \n để xuống dòng
        txt_thanh_vien += f"\n\t◦ {mem['ten']}\t\tĐơn vị: {mem['don_vi']}"

    # Thời gian kết thúc (Chỉ giờ phút)
    txt_time_end = "11:30"

    # Kết quả phiếu
    dong_y = 0      # 0 phiếu thuận
    phan_doi = 5    # 5 phiếu chống
    is_approved = False # REJECTED

    context = {
        "ten_khoa": "CÔNG NGHỆ THÔNG TIN",
        "QD_so": "1001/QĐ-ĐHSPKT",
        "ngay_cua_hieu_truong": "15/01/2024",
        "thoi_gian_hop": "08:00 ngày 20/01/2024",
        "thoi_gian_ket_thuc": txt_time_end, # Chỉ giờ
        "dia_diem": "Phòng họp số 2",
        "ten_de_tai": "Nghiên cứu xây dựng hệ thống quản lý văn bản",
        
        # Map biến vào Template
        "chu_nhiem": txt_chu_nhiem, 
        "thanh_vien": txt_thanh_vien, 
        
        "co_mat_tren_tong_so": "05/05",
        "vang_mat": "0",
        "so_phieu_phat_ra": "5", "so_phieu_thu_vao": "5",
        "so_phieu_dong_y": "5", "so_phieu_phan_doi": "0",
        "box_de_nghi": "[x]" if is_approved else "[ ]",
        "box_khong_de_nghi": "[ ]" if is_approved else "[x]",
        
        "nhung_noi_dung_chinh_sua": "- Bổ sung tài liệu." if is_approved else "",
        "nhung_noi_dung_bo_sung": "- Thêm module báo cáo." if is_approved else "",
        
        # Nếu ĐẠT -> Để rỗng biến này (không điền "Không có.")
        "nhung_noi_dung_khong_phu_hop_sau": "" if is_approved else "- Tính cấp thiết chưa rõ.",
        
        "ten_chu_tich": "PGS.TS. Lê Văn Chủ Tịch",
        "ten_thu_ky": "ThS. Nguyễn Văn Thư Ký"
    }

    # 3. MANUAL RENDER & CLEAN UP
    tmpl_path = "form_engine/templates/3b.docx"
    out_docx = "form_engine/output/2026-01-16/3b_rejected.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()
    
    # Replace
    for p in doc.paragraphs: engine.replace_text_in_element(p, context)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: engine.replace_text_in_element(p, context)

    # Clean Up Logic (Đã test ổn)
    paragraphs_to_delete = []
    KEY_POSITIVE = "Đề nghị Nhà trường cho phép thực hiện"
    KEY_NEGATIVE = "Đề nghị Nhà trường không cho phép thực hiện"
    KEY_OR = "Hoặc"
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == KEY_OR:
            paragraphs_to_delete.append(p); continue
        if is_approved:
            if KEY_NEGATIVE in text or "không phù hợp sau:" in text: paragraphs_to_delete.append(p)
        else:
            if KEY_POSITIVE in text or "chỉnh sửa, bổ sung" in text: paragraphs_to_delete.append(p)

    for p in paragraphs_to_delete:
        if p._element.getparent() is not None: p._element.getparent().remove(p._element)

    doc.save(out_docx)
    
    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created Final 3b: {out_docx.replace('.docx', '.pdf')}")
    except: pass

if __name__ == "__main__":
    main()
