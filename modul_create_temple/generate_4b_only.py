from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_format(cell, text, bold=False, align=None):
    cell.text = text
    # Set Font & Format
    for paragraph in cell.paragraphs:
        if align: paragraph.alignment = align
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13) # Size 13 theo yêu cầu
            run.bold = bold

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Hack to set cell borders directly in XML
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create tcBorders element
    tcBorders = OxmlElement('w:tcBorders')
    
    # Define borders
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        element = OxmlElement(f'w:{border}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4') # Size 4 = 0.5pt
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000') # Black
        tcBorders.append(element)
        
    tcPr.append(tcBorders)

def main():
    print("--- GENERATING FORM 4B (FORMATTED + BORDER FIX) ---")
    
    # ... (Data)
    # 1. Prepare Data
    table_items = [
        {"stt": 1, "ten": "Nghiên cứu xây dựng Chatbot AI tư vấn tuyển sinh", "cn": "TS. Nguyễn Văn A", "mt": "Tự động hóa 90% câu trả lời", "tm": "Mới", "kq": "Phần mềm + Báo cáo", "ud": "Phòng Đào tạo", "kp": "15.000.000"},
        {"stt": 2, "ten": "Hệ thống điểm danh sinh viên bằng khuôn mặt", "cn": "ThS. Trần Thị B", "mt": "Chính xác 99%", "tm": "Cải tiến", "kq": "Thiết bị + App", "ud": "Giảng đường", "kp": "20.000.000"},
        {"stt": 3, "ten": "Ứng dụng Blockchain trong quản lý văn bằng", "cn": "KS. Lê Văn C", "mt": "Chống làm giả", "tm": "Mới", "kq": "Website tra cứu", "ud": "Phòng CTSV", "kp": "12.000.000"},
        {"stt": 4, "ten": "Phân tích dữ liệu học tập Big Data", "cn": "TS. Phạm Văn D", "mt": "Dự báo sinh viên bỏ học", "tm": "Sáng tạo", "kq": "Báo cáo phân tích", "ud": "Ban Giám hiệu", "kp": "10.000.000"},
        {"stt": 5, "ten": "Xây dựng Lab thực hành ảo Cloud Computing", "cn": "ThS. Hoàng Văn E", "mt": "Tiết kiệm chi phí phần cứng", "tm": "Cải tiến", "kq": "Hệ thống Lab Online", "ud": "Sinh viên CNTT", "kp": "30.000.000"}
    ]

    context = {
        "ten_khoa": "Công nghệ Thông tin".upper(), # Tự động IN HOA
        "nam_hoc": "2024 - 2025",
        "ten_chu_tich": "PGS.TS. Lê Văn Trưởng Khoa",
        "ten_thu_ky": "ThS. Phạm Thị Thư Ký",
        "ngay": "20", "thang": "01", "nam": "2024",
        "so_de_tai": str(len(table_items)) # Tự động đếm: 5
    }
    
    tmpl_path = "form_engine/templates/4b.docx"
    out_docx = "form_engine/output/2026-01-15/4b_formatted.docx"
    
    doc = Document(tmpl_path)
    
    # A. Fill Variables (FULL SCAN: Body, Tables, Header, Footer)
    
    # 1. Body Paragraphs
    for p in doc.paragraphs:
        for k, v in context.items():
            if f"{{{{ {k} }}}}" in p.text: p.text = p.text.replace(f"{{{{ {k} }}}}", str(v))
            if f"{{{{{k}}}}}" in p.text: p.text = p.text.replace(f"{{{{{k}}}}}", str(v))

    # 2. All Tables (Header Table, Footer Table, Data Table)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in context.items():
                        if f"{{{{ {k} }}}}" in p.text: p.text = p.text.replace(f"{{{{ {k} }}}}", str(v))
                        if f"{{{{{k}}}}}" in p.text: p.text = p.text.replace(f"{{{{{k}}}}}", str(v))

    # 3. Section Headers & Footers (Nếu biến nằm trong Header/Footer chuẩn của Word)
    for section in doc.sections:
        # Check Header
        for p in section.header.paragraphs:
            for k, v in context.items():
                if f"{{{{ {k} }}}}" in p.text: p.text = p.text.replace(f"{{{{ {k} }}}}", str(v))
                if f"{{{{{k}}}}}" in p.text: p.text = p.text.replace(f"{{{{{k}}}}}", str(v))
        # Check Footer
        for p in section.footer.paragraphs:
            for k, v in context.items():
                if f"{{{{ {k} }}}}" in p.text: p.text = p.text.replace(f"{{{{ {k} }}}}", str(v))
                if f"{{{{{k}}}}}" in p.text: p.text = p.text.replace(f"{{{{{k}}}}}", str(v))
        # Check Tables in Header/Footer (Hiếm nhưng có thể)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k, v in context.items():
                            if f"{{{{ {k} }}}}" in p.text: p.text = p.text.replace(f"{{{{ {k} }}}}", str(v))
                            if f"{{{{{k}}}}}" in p.text: p.text = p.text.replace(f"{{{{{k}}}}}", str(v))

    # B. Fill Table
    if len(doc.tables) > 1:
        target_table = doc.tables[1]
        
        # BỎ DÒNG SET STYLE GÂY LỖI
        # target_table.style = 'Table Grid'
        
        print(f"   Injecting data into Table 1 ({len(target_table.columns)} cols)...")
        
        for item in table_items:
            row = target_table.add_row()
            cells = row.cells
            
            # Map data & Format & Border
            # [0] STT
            set_cell_format(cells[0], str(item['stt']), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_border(cells[0])
            
            # [1] Tên
            set_cell_format(cells[1], item['ten'])
            set_cell_border(cells[1])
            
            # [2] BCN
            set_cell_format(cells[2], item['cn'])
            set_cell_border(cells[2])
            
            # [3] Mục tiêu
            if len(cells) > 3: 
                set_cell_format(cells[3], item['mt'])
                set_cell_border(cells[3])
            
            # [4] Tính mới
            if len(cells) > 4: 
                set_cell_format(cells[4], item['tm'])
                set_cell_border(cells[4])
            
            # [5] Nội dung
            if len(cells) > 5: 
                set_cell_format(cells[5], "- Nội dung 1\n- Nội dung 2")
                set_cell_border(cells[5])
            
            # [6] Kết quả
            if len(cells) > 6: 
                set_cell_format(cells[6], item['kq'])
                set_cell_border(cells[6])
            
            # [7] Ứng dụng
            if len(cells) > 7: 
                set_cell_format(cells[7], item['ud'])
                set_cell_border(cells[7])
            
            # [8] Kinh phí
            if len(cells) > 8: 
                set_cell_format(cells[8], item['kp'], align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_border(cells[8])

        # Xóa dòng mẫu
        if len(target_table.rows) > 1:
            row_1_text = target_table.rows[1].cells[0].text
            if "{{" in row_1_text:
                 tbl = target_table._tbl
                 tr = target_table.rows[1]._tr
                 tbl.remove(tr)
    
    doc.save(out_docx)
    
    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-15", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created Formatted 4b: {out_docx.replace('.docx', '.pdf')}")
    except: pass

if __name__ == "__main__":
    main()
