import sys
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine

def main():
    print("--- GENERATING FORM 5B (XÉT CHỌN SƠ BỘ) ---")
    
    # 1. Prepare Data
    context = {
        # Header
        "so": "10/BB-ĐHSPKT",
        "qd_so": "123/QĐ-ĐHSPKT", # <--- ĐÃ BỔ SUNG
        "thoi_gian_hop": "08:00 ngày 25/01/2024",
        "dia_diem_hop": "Phòng họp Hiệu bộ",
        
        # Thành phần
        "ten_nguoi_chu_tri": "PGS.TS. Lê Hiệu Trưởng", # Chủ tọa
        "ten_thu_ky": "TS. Phạm Trưởng Phòng KHCN",
        "co_mat_tren_tong": "07/07",
        "vang_mat": "0",
        
        # Nội dung (Text dài có xuống dòng)
        "tien_trinh_hop": (
            "1. Chủ tịch Hội đồng tuyên bố lý do, giới thiệu đại biểu.\n"
            "2. Các chủ nhiệm đề tài trình bày tóm tắt thuyết minh (5 phút/đề tài).\n"
            "3. Hội đồng thảo luận, đánh giá tính khả thi của từng đề tài.\n"
            "4. Hội đồng bỏ phiếu xét chọn sơ bộ."
        ),
        
        "noi_dung_hop": (
            "- Đánh giá cao tính cấp thiết của các đề tài.\n"
            "- Yêu cầu chỉnh sửa lại tên đề tài số 3 cho gọn hơn.\n"
            "- Đề nghị bổ sung thêm thành viên nghiên cứu cho đề tài số 5."
        ),
        
        # Kết luận
        "ket_luan_cua_hoi_dong": (
            "Hội đồng thống nhất đưa 05 đề tài vào danh sách xét chọn chính thức cấp Trường.\n"
            "Giao Phòng KHCN hướng dẫn các chủ nhiệm hoàn thiện hồ sơ."
        ),
        
        # Footer
        "ten_chu_toa": "PGS.TS. Lê Hiệu Trưởng",
        "ngay_ht": "25", "thang": "01", "nam": "2024"
    }

    # 2. Render
    tmpl_path = "form_engine/templates/5b.docx"
    out_docx = "form_engine/output/2026-01-16/5b_generated.docx"
    
    # Sử dụng Engine Wrapper (có sẵn hàm Smart Replace và Newline Support)
    engine = FormEngine()
    
    # Do logic 5b đơn giản (không có checkbox phức tạp hay xóa dòng), 
    # ta có thể dùng engine.render() trực tiếp nếu template chuẩn.
    # Nhưng để chắc chắn (vì template có thể cần xử lý Header/Footer), ta dùng manual loop.
    
    doc = Document(tmpl_path)
    
    # Replace Everywhere with ALIGNMENT FIX
    # Danh sách các biến cần căn trái (Nội dung dài)
    vars_need_left_align = ["tien_trinh_hop", "noi_dung_hop", "ket_luan_cua_hoi_dong"]
    
    def process_paragraph(p, ctx):
        # 1. Check if alignment fix needed
        for var in vars_need_left_align:
            if f"{{{{ {var} }}}}" in p.text or f"{{{{{var}}}}}" in p.text:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 2. Replace text
        engine.replace_text_in_element(p, ctx)

    for p in doc.paragraphs: process_paragraph(p, context)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: process_paragraph(p, context)
                
    # Check Header/Footer (Quan trọng cho 5b vì nó là biên bản chính thức)
    for section in doc.sections:
        for p in section.header.paragraphs: engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs: engine.replace_text_in_element(p, context)

    # Save
    doc.save(out_docx)
    
    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 5b: {out_docx.replace('.docx', '.pdf')}")
    except: pass

if __name__ == "__main__":
    main()
