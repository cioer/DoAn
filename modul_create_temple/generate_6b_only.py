import sys
import os
from docx import Document

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine

def main():
    print("--- GENERATING FORM 6B (COPY LOGIC FROM 3B) ---")
    
    is_approved = False # REJECTED
    
    context = {
        "qd_so": "123/QĐ-ĐHSPKT", "qd_sp": "123/QĐ-ĐHSPKT", "QD_so": "123/QĐ-ĐHSPKT",
        "thoi_gian_hop": "14:00 ngày 26/01/2024",
        "dia_diem": "Phòng họp Khoa CNTT",
        "ten_nguoi_chu_tri": "TS. Trưởng Khoa", 
        "ten_thu_ky": "ThS. Thư Ký Khoa",
        "co_mat_tren_tong": "05/05",
        "vang_mat": "0",
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "nam_hoc": "2024-2025",
        "thoi_gian_ket_thuc": "16:30",
        "ngay_ht": "26", "thang": "01", "nam": "2024",
        
        # Logic nội dung (Simple string replacement)
        "noi_dung_da_chinh_sua": ("- Đã cập nhật lại tên đề tài.\n- Đã bổ sung phương pháp.") if is_approved else " ",
        "noi_dung_bo_sung": ("- Thêm thành viên nghiên cứu.") if is_approved else " ",
        "noi_dung_khong_phu_hop": (" ") if is_approved else ("- Đề cương sơ sài, không đạt yêu cầu khoa học."),
    }

    tmpl_path = "form_engine/templates/6b.docx"
    out_docx = "form_engine/output/2026-01-16/6b_rejected_simple.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()
    
    # 1. Replace Simple (Như 3b)
    for p in doc.paragraphs: engine.replace_text_in_element(p, context)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: engine.replace_text_in_element(p, context)
    for section in doc.sections:
        for p in section.header.paragraphs: engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs: engine.replace_text_in_element(p, context)

    # 2. Clean Up Logic
    paragraphs_to_delete = []
    KEY_POSITIVE = "Đề nghị Nhà trường cho phép thực hiện"
    KEY_NEGATIVE = "Đề nghị Nhà trường không cho phép thực hiện"
    KEY_OR = "Hoặc"
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == KEY_OR: paragraphs_to_delete.append(p); continue
        if is_approved:
            if KEY_NEGATIVE in text or "nội dung không phù hợp sau" in text: paragraphs_to_delete.append(p)
        else:
            if KEY_POSITIVE in text or "nội dung chỉnh sửa" in text or "nội dung bổ sung" in text: paragraphs_to_delete.append(p)

    for p in paragraphs_to_delete:
        if p._element.getparent() is not None: p._element.getparent().remove(p._element)

    # 3. Clean Empty Lines
    for i in range(len(doc.paragraphs) - 1, 0, -1):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if not text or text in [".", ",", "_"]:
             try: p._element.getparent().remove(p._element)
             except: pass
        if "}}" in p.text: p.text = p.text.replace("}}", "")

    doc.save(out_docx)
    
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 6b Simple: {out_docx.replace('.docx', '.pdf')}")
    except: pass

if __name__ == "__main__":
    main()
