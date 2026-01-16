import sys
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine

def main():
    print("--- GENERATING FORM 7B (BÁO CÁO HOÀN THIỆN) ---")
    
    # 1. Prepare Data
    # Helper tạo ngày tháng không ngắt dòng
    def make_date_line(d, m, y):
        # Sử dụng Non-breaking space (\u00A0) thay cho space thường
        return f"Ngày\u00A0{d}\u00A0tháng\u00A0{m}\u00A0năm\u00A0{y}"

    date_str = make_date_line("28", "01", "2024")

    context = {
        # Header
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "nguoi_de_xuat": "TS. Nguyễn Văn AI",
        "ten_nguoi_de_xuat": "TS. Nguyễn Văn AI", # Biến thể tên
        "thoi_gian_hop": "14:00 ngày 26/01/2024",

        # Hội đồng (THIẾU TRONG BẢN GỐC)
        "ten_chu_tich_hoi_dong": "PGS.TS. Trần Văn B",
        "ten_thu_ky": "ThS. Lê Thị C",

        # Nội dung chính (Cần căn trái)
        "noi_dung_da_sua": (
            "- Đã chỉnh sửa lại tên đề tài theo kết luận của Hội đồng.\n"
            "- Đã làm rõ hơn phương pháp nghiên cứu định lượng."
        ),

        "noi_dung_bo_sung": (
            "- Đã bổ sung 01 thành viên nghiên cứu là sinh viên.\n"
            "- Đã thêm kế hoạch khảo sát thực tế."
        ),

        # Ngày tháng
        "ngay": "28", "thang": "01", "nam": "2024"
    }
    
    # [NEW STRATEGY]
    # Nếu trong file Word đang là: "Nam Định, ngày {{ngay}} tháng {{thang}} năm {{nam}}"
    # Ta sẽ tìm và thay thế CẢ CỤM đó bằng chuỗi Non-breaking.
    
    special_date_replace = True
    location_prefix = "Nam Định, " # Hoặc lấy từ config

    # 2. Render
    tmpl_path = "form_engine/templates/7b.docx"
    out_docx = "form_engine/output/2026-01-16/7b_clean.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()
    
    # Danh sách biến cần căn trái
    vars_left_align = ["noi_dung_da_sua", "noi_dung_bo_sung"]

    def process_element(p):
        # 1. Date Line Hard Fix (Xử lý dòng ngày tháng)
        # Tìm dòng chứa pattern ngày tháng năm
        if "ngày" in p.text and "tháng" in p.text and "năm" in p.text and "{{" in p.text:
            # Thay thế toàn bộ dòng bằng chuỗi dính liền
            # Giả sử format gốc: "..., ngày {{ngay}} tháng {{thang}} năm {{nam}}"
            # Ta sẽ thay thế các khoảng trắng quanh nó bằng \u00A0
            
            # Cách đơn giản: Replace text bình thường trước
            engine.replace_text_in_element(p, context)
            
            # Sau đó replace space -> non-breaking space cho dòng này
            # (Lưu ý: Chỉ làm nếu dòng này ngắn, ví dụ < 100 ký tự, để tránh dính cả đoạn văn dài)
            if len(p.text) < 100:
                p.text = p.text.replace(" ", "\u00A0")
                return # Done for this paragraph

        # 2. Alignment Fix
        for var in vars_left_align:
            if var in p.text:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 3. Replace normal
        engine.replace_text_in_element(p, context)

    # Loop Replace
    for p in doc.paragraphs: process_element(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: process_element(p)
    
    # Header/Footer
    for section in doc.sections:
        for p in section.header.paragraphs: process_element(p)
        for p in section.footer.paragraphs: process_element(p)

    # Save
    doc.save(out_docx)
    
    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 7b: {out_docx.replace('.docx', '.pdf')}")
    except: pass

if __name__ == "__main__":
    main()
