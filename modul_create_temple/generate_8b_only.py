import sys
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine

def main():
    print("--- GENERATING FORM 8B (ĐỀ NGHỊ NGHIỆM THU) ---")
    
    # 1. Prepare Data
    # Danh sách ỦY VIÊN (Chủ tịch và Thư ký đã điền ở row 1, 2)
    # STT bắt đầu từ 3 vì: 1=Chủ tịch, 2=Thư ký
    hoi_dong = [
        {"stt": 3, "ten": "TS. Phạm Văn C", "don_vi": "Khoa Điện"},
        {"stt": 4, "ten": "TS. Hoàng Văn D", "don_vi": "Khoa Cơ khí"},
        {"stt": 5, "ten": "ThS. Ngô Thị E", "don_vi": "Trung tâm Số"}
    ]

    context = {
        "ten_khoa": "CÔNG NGHỆ THÔNG TIN",
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "ma_de_tai": "NCKH-2024-01",

        # Ban chủ nhiệm
        "ten_chu_nhiem": "TS. Nguyễn Văn AI",

        # Chữ ký Hội đồng nghiệm thu
        "ten_chu_tich": "PGS.TS. Trần Văn B",
        "don_vi_cong_tac_chu_tich": "Khoa CNTT - Trường ĐH SPKT",

        "ten_thu_ky": "ThS. Lê Thị C",
        "don_vi_cong_tac_thu_ky": "Khoa CNTT - Trường ĐH SPKT",

        # Date Line (Quy chuẩn mới)
        "diadiem_thoigian": "Nam Định,\u00A0ngày\u00A030\u00A0tháng\u00A001\u00A0năm\u00A02024",
        "ngay": "30", "thang": "01", "nam": "2024" # Fallback cũ
    }

    # 2. Render
    tmpl_path = "form_engine/templates/8b.docx"
    out_docx = "form_engine/output/2026-01-16/8b_generated.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()

    # Replace Variables in ALL places (paragraphs + tables + headers + footers)
    for p in doc.paragraphs:
        engine.replace_text_in_element(p, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    engine.replace_text_in_element(p, context)

    for section in doc.sections:
        for p in section.header.paragraphs:
            engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs:
            engine.replace_text_in_element(p, context)

    # 3. FILL TABLE HỘI ĐỒNG (Bảng thứ 2 - Index 1)
    # Bảng này có cấu trúc: Row 0 = Header, Row 1 = Chủ tịch, Row 2 = Thư ký, Row 3 = Template ủy viên
    if len(doc.tables) > 1:
        target_table = doc.tables[1]
        print(f"   Found council table with {len(target_table.rows)} rows. Filling members...")

        # Helper function để điền text vào cell KHÔNG in đậm
        def fill_cell_text(cell, text, align=None):
            """Điền text vào cell và reset formatting (bỏ bold)"""
            # Xóa tất cả runs cũ
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ""

            # Tạo run mới với text và format bình thường
            if cell.paragraphs:
                p = cell.paragraphs[0]
                p.text = text
                # Reset formatting: không bold
                for run in p.runs:
                    run.bold = False
                    run.font.name = 'Times New Roman'

                # Căn lề
                if align:
                    p.alignment = align

        # Bắt đầu từ row 3 (template ủy viên), điền đè lên
        start_row_idx = 3

        # Xử lý riêng: Chủ tịch (row 1) và Thư ký (row 2) đã được replace ở trên
        # Bây giờ chỉ cần thêm các ủy viên
        for i, mem in enumerate(hoi_dong):
            row_idx = start_row_idx + i

            # Nếu bảng thiếu dòng -> Thêm dòng
            if row_idx >= len(target_table.rows):
                row = target_table.add_row()
            else:
                row = target_table.rows[row_idx]

            cells = row.cells
            # Map cột: [0] STT | [1] Họ tên | [2] Đơn vị công tác | [3] Chức danh
            if len(cells) >= 4:
                fill_cell_text(cells[0], str(mem['stt']), WD_ALIGN_PARAGRAPH.CENTER)
                fill_cell_text(cells[1], mem['ten'], WD_ALIGN_PARAGRAPH.LEFT)
                fill_cell_text(cells[2], mem['don_vi'], WD_ALIGN_PARAGRAPH.LEFT)
                fill_cell_text(cells[3], "Ủy viên", WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    doc.save(out_docx)
    
    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 8b: {out_docx.replace('.docx', '.pdf')}")
    except: pass

if __name__ == "__main__":
    main()
