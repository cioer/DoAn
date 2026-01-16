import sys
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine, set_left_align_for_lists, CHECKBOX_CHECKED, CHECKBOX_UNCHECKED

def main():
    print("--- GENERATING FORM 9B (PHIẾU ĐÁNH GIÁ NGHIỆM THU) ---")

    # 1. Prepare Data
    context = {
        # Thông tin đề tài
        "ma_de_tai": "NCKH-2024-01",
        "ten_de_tai": "Nghiên cứu ứng dụng Blockchain trong quản lý văn bằng",
        "ten_chu_nhiem": "TS. Nguyễn Văn AI",

        # Thông tin hội đồng
        "don_vi_chu_tri": "Khoa Công nghệ thông tin",
        "don_vi_cong_tac": "Trường ĐH Sư phạm Kỹ thuật Nam Định",

        # Thời gian, địa điểm
        "thoi_gian_hop": "14:00 ngày 30/01/2024",
        "dia_diem_hop": "Phòng họp Khoa CNTT",

        # Người đánh giá
        "ten_nguoi_danh_gia": "PGS.TS. Trần Văn B",
        "ten_thanh_vien": "TS. Phạm Văn C, TS. Hoàng Văn D",

        # Nội dung đánh giá - Các chỉ tiêu khoa học
        "cac_chi_tieu_chu_yeu_cac_yeu_cau_khoa_hoc_cua_ket_qua": (
            "- Đề tài đáp ứng các chỉ tiêu khoa học cơ bản.\n"
            "- Phương pháp nghiên cứu phù hợp, kết quả có giá trị thực tiễn.\n"
            "- Số lượng bài báo khoa học: 02 bài đăng trên tạp chí quốc tế.\n"
            "- Sản phẩm nghiệm thu: Hệ thống quản lý văn bằng demo."
        ),

        # Phương pháp nghiên cứu
        "phuong_phap_nghien_cuu": (
            "Kết hợp nghiên cứu lý thuyết và thực nghiệm:\n"
            "- Khảo sát thực trạng tại 3 đơn vị đại học.\n"
            "- Thiết kế và xây dựng hệ thống Blockchain.\n"
            "- Kiểm thử và đánh giá hiệu quả."
        ),

        # Số lượng chung, loại, khối lượng sản phẩm
        "so_luong_chung_loai_khoi_luong_san_pham": (
            "- 01 hệ thống quản lý văn bằng trên nền Blockchain.\n"
            "- 02 bài báo khoa học (01 quốc tế, 01 trong nước).\n"
            "- 01 báo cáo tổng hợp độ dài 50 trang."
        ),

        # Nhận xét về mức độ hoàn thành
        "nhan_xet_ve_muc_do_hoan_thanh": (
            "Đề tài đã hoàn thành đầy đủ các nội dung theo đề xuất ban đầu. "
            "Kết quả nghiên cứu có tính ứng dụng cao, có thể triển khai thực tế. "
            "Chủ nhiệm đề tài và nhóm thực hiện đầy đủ trách nhiệm."
        ),

        # Checkbox - QUY TẮC CHUNG: Dùng [x]/[ ]
        "box_dat": CHECKBOX_CHECKED,
        "box_khong_dat": CHECKBOX_UNCHECKED,

        # Ý kiến khác
        "y_kien_khac": "Đề nghị chủ nhiệm đề tài tiếp tục phát triển và mở rộng hệ thống."
    }

    # 2. Render
    tmpl_path = "form_engine/templates/9b.docx"
    out_docx = "form_engine/output/2026-01-16/9b_generated.docx"
    doc = Document(tmpl_path)
    engine = FormEngine()

    # Replace Variables in ALL places
    for p in doc.paragraphs:
        engine.replace_text_in_element(p, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    engine.replace_text_in_element(p, context)

    for section in doc.sections:
        for p in section.header.paragraphs:
            engine.replace_text_in_element(p, context)
        for p in section.footer.paragraphs:
            engine.replace_text_in_element(p, context)

    # 3. Căn trái các danh sách gạch đầu dòng (QUY TẸC CHUNG)
    # Tự động nhận diện dòng bắt đầu bằng '-' hoặc '•'
    set_left_align_for_lists(doc)

    # Save
    doc.save(out_docx)

    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-16", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created 9b: {out_docx.replace('.docx', '.pdf')}")
    except:
        print(f"⚠️ PDF conversion failed")

    # Verify unfilled variables
    doc_check = Document(out_docx)
    import re
    text = ""
    for p in doc_check.paragraphs:
        text += p.text + "\n"
    for table in doc_check.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text += p.text + "\n"

    unfilled = re.findall(r'\{\{[^}]+\}\}', text)
    if unfilled:
        print(f"⚠️ Còn {len(set(unfilled))} biến chưa điền:")
        for u in sorted(set(unfilled)):
            print(f"     {u}")
    else:
        print("✅ Tất cả biến đã điền đầy đủ!")

if __name__ == "__main__":
    main()
