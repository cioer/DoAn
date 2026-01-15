import sys
import os
import glob
import logging

# Add path
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))

from form_engine.src.core.engine import FormEngine
from form_engine.src.schemas.all_forms import Form1b, Form2b, Form3b, Form4b, GenericForm

# Setup
engine = FormEngine()
logging.basicConfig(level=logging.INFO)

def get_template_list():
    path = "form_engine/templates/*.docx"
    files = glob.glob(path)
    return [os.path.basename(f) for f in files]

def generate_form_4b(filename):
    """Xử lý riêng cho bảng danh sách"""
    print(f"   [Special] Processing Table for {filename}...")
    
    # Dữ liệu bảng
    items = [
        {"stt": 1, "ten": "Nghiên cứu AI", "cn": "Nguyễn Văn A", "kp": "10tr"},
        {"stt": 2, "ten": "Nghiên cứu IoT", "cn": "Trần Văn B", "kp": "15tr"},
        {"stt": 3, "ten": "Nghiên cứu BigData", "cn": "Lê Văn C", "kp": "20tr"},
    ]
    
    # Input data cơ bản
    data = Form4b(ten_khoa="CNTT", nam_hoc="2024-2025").dict()
    
    # Vì Engine hiện tại là "Simple Replace", ta không thể inject row vào bảng 
    # trừ khi ta viết code thao tác bảng trực tiếp ở đây.
    # Để demo nhanh, tôi sẽ dùng cách inject trực tiếp vào docx object rồi save.
    
    from docx import Document
    template_path = os.path.join("form_engine/templates", filename)
    doc = Document(template_path)
    
    # 1. Fill biến thường (dùng engine helper logic - tái sử dụng code engine một chút)
    # Tuy nhiên để nhanh, ta replace thủ công ở đây
    for p in doc.paragraphs:
        for k, v in data.items():
            if f"{{{{ {k} }}}}" in p.text:
                p.text = p.text.replace(f"{{{{ {k} }}}}", str(v))
    
    # 2. Fill Table (Tìm bảng đầu tiên)
    if len(doc.tables) > 0:
        table = doc.tables[0]
        for item in items:
            row = table.add_row().cells
            # Giả định bảng có 5 cột
            if len(row) >= 5:
                row[0].text = str(item['stt'])
                row[1].text = item['ten']
                row[2].text = item['cn']
                # Cột 3, 4 tùy mẫu
                row[4].text = item['kp']
    
    # Save & Convert
    out_docx = f"form_engine/output/2026-01-15/{filename.replace('.docx', '_filled.docx')}"
    doc.save(out_docx)
    
    # Convert PDF call
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-15", out_docx]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"   ✅ OK: {out_docx}")
    except:
        print("   ⚠️ PDF Failed")

def main():
    templates = get_template_list()
    templates.sort()
    
    print(f"--- BATCH GENERATING {len(templates)} FORMS ---\n")
    
    for tmpl in templates:
        print(f"🚀 Processing: {tmpl}")
        
        try:
            # 1. Router: Chọn Schema phù hợp
            if "4b" in tmpl:
                generate_form_4b(tmpl)
                continue
                
            elif "1b" in tmpl:
                data = Form1b(
                    tinh_cap_thiet="Rất cấp thiết (Demo Batch)",
                    noi_dung_chinh="- Nội dung 1\n- Nội dung 2"
                ).dict()
                
            elif "2b" in tmpl:
                data = Form2b(ket_qua_dat=True).dict()
                
            elif "3b" in tmpl:
                # Demo list thành viên
                members = "- Nguyễn Văn A (Thư ký)\n- Trần Văn B (Ủy viên)"
                data = Form3b(danh_sach_thanh_vien=members, thanh_vien=members).dict()
                
            else:
                # Các form còn lại dùng Generic
                data = GenericForm(
                    is_pass=True,
                    thanh_vien="- Nguyễn Văn A\n- Trần Văn B", # Fallback cho các form có ds thành viên
                    nhung_thanh_vien="- Nguyễn Văn A\n- Trần Văn B",
                    qd_sp="123/QĐ-SP" # Fix typo support
                ).dict()
            
            # 2. Render
            res = engine.render(tmpl, data, user_id="batch_job")
            print(f"   ✅ OK: {res['pdf']}")
            
        except Exception as e:
            print(f"   ❌ FAILED: {e}")

if __name__ == "__main__":
    main()
