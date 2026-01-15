import re
from docx import Document
import os
import glob

def scan_all_templates():
    template_dir = 'form_engine/templates'
    files = glob.glob(os.path.join(template_dir, "*.docx"))
    files.sort()
    
    print(f"--- SCANNING {len(files)} TEMPLATES ---")
    
    for path in files:
        filename = os.path.basename(path)
        try:
            doc = Document(path)
            text = ""
            # Scan Paragraphs
            for p in doc.paragraphs:
                text += p.text + "\n"
            # Scan Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
                        # Deep scan nested paragraphs in cells
                        for p in cell.paragraphs:
                            text += p.text + "\n"
            
            # Regex tìm {{ variable }}
            tags = set(re.findall(r'\{\{\s*(\w+)\s*\}\}', text))
            
            if tags:
                print(f"\n📄 {filename}:")
                for t in sorted(list(tags)):
                    print(f"   - {t}")
            else:
                print(f"\n⚪ {filename}: (Không tìm thấy biến nào)")
                
        except Exception as e:
            print(f"\n❌ {filename}: Lỗi đọc file ({e})")

if __name__ == "__main__":
    scan_all_templates()
