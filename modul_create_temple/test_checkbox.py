from docx import Document
import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine

# 1. TẠO TEMPLATE MẪU CHECKBOX
def create_checkbox_template():
    doc = Document()
    doc.add_heading('PHIẾU KHẢO SÁT (DEMO CHECKBOX)', 0)
    
    p = doc.add_paragraph()
    p.add_run('Họ và tên: Nguyễn Văn A').bold = True
    
    # Checkbox Giới tính
    p = doc.add_paragraph()
    p.add_run('Giới tính: ')
    p.add_run('{{ box_nam }} Nam      ')
    p.add_run('{{ box_nu }} Nữ')
    
    # Checkbox Sở thích (Chọn nhiều)
    p = doc.add_paragraph()
    p.add_run('Sở thích: ')
    p.add_run('{{ box_bong_da }} Bóng đá    ')
    p.add_run('{{ box_doc_sach }} Đọc sách    ')
    p.add_run('{{ box_du_lich }} Du lịch')

    # Checkbox Đánh giá (Đạt/Không)
    p = doc.add_paragraph()
    p.add_run('Kết luận: ')
    p.add_run('{{ box_dat }} Đạt      ')
    p.add_run('{{ box_khong_dat }} Không đạt')
    
    path = 'form_engine/templates/test_checkbox.docx'
    doc.save(path)
    print(f"Created template: {path}")

# 2. LOGIC ĐIỀN DỮ LIỆU
def run_checkbox_test():
    engine = FormEngine()
    
    # Dữ liệu giả định
    input_gender = "NAM"  # hoặc "NU"
    input_hobbies = ["BONG_DA", "DU_LICH"] # Đọc sách không chọn
    input_result = False # Không đạt
    
    # Ký tự Checkbox (Bạn có thể copy paste các ký tự này)
    CHECKED = "☑"
    UNCHECKED = "☐"
    
    # Logic xử lý biến
    data = {}
    
    # 1. Xử lý Giới tính
    if input_gender == "NAM":
        data['box_nam'] = CHECKED
        data['box_nu'] = UNCHECKED
    else:
        data['box_nam'] = UNCHECKED
        data['box_nu'] = CHECKED
        
    # 2. Xử lý Sở thích (Đa lựa chọn)
    data['box_bong_da'] = CHECKED if "BONG_DA" in input_hobbies else UNCHECKED
    data['box_doc_sach'] = CHECKED if "DOC_SACH" in input_hobbies else UNCHECKED
    data['box_du_lich'] = CHECKED if "DU_LICH" in input_hobbies else UNCHECKED
    
    # 3. Xử lý Kết quả
    if input_result: # True là Đạt
        data['box_dat'] = CHECKED
        data['box_khong_dat'] = UNCHECKED
    else:
        data['box_dat'] = UNCHECKED
        data['box_khong_dat'] = CHECKED # Đánh dấu vào ô Không Đạt
        
    print("Data prepared:", data)
    
    res = engine.render("test_checkbox.docx", data, user_id="test_checkbox")
    print(f"✅ Result: {res['pdf']}")

if __name__ == "__main__":
    create_checkbox_template()
    run_checkbox_test()
