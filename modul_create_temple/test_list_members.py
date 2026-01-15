from docx import Document
import os
import sys
sys.path.append(os.path.join(os.path.dirname(__file__), 'form_engine/src'))
from form_engine.src.core.engine import FormEngine

# 1. TẠO TEMPLATE 3B
def create_template_3b():
    doc = Document()
    doc.add_heading('MẪU 3B - KIỂM TRA DANH SÁCH THÀNH VIÊN', 0)
    
    p = doc.add_paragraph()
    p.add_run('1. Chủ nhiệm đề tài: ').bold = True
    p.add_run('{{ chu_nhiem }}')
    
    p = doc.add_paragraph()
    p.add_run('2. Danh sách thành viên tham gia:').bold = True
    
    # Biến này sẽ chứa nhiều dòng
    p = doc.add_paragraph('{{ danh_sach_thanh_vien }}')
    
    path = 'form_engine/templates/3b.docx'
    doc.save(path)
    print(f"Created {path}")

# 2. CHẠY THỬ
def run_test():
    engine = FormEngine()
    
    # Giả lập danh sách lấy từ DB
    members_db = [
        {"name": "Nguyễn Văn A", "role": "Thư ký"},
        {"name": "Trần Thị B",   "role": "Ủy viên"},
        {"name": "Lê Văn C",     "role": "Kỹ thuật viên"},
        {"name": "Phạm Thị D",   "role": "Phân tích dữ liệu"}
    ]
    
    # Logic nối chuỗi (Pre-processing)
    # Tạo chuỗi: "- Tên (Vai trò)" và xuống dòng
    members_str = "\n".join([f"- {m['name']} ({m['role']})" for m in members_db])
    
    data = {
        "chu_nhiem": "GS.TS. Ngô Bảo Châu",
        "danh_sach_thanh_vien": members_str
    }
    
    print(f"Data Members:\n{members_str}")
    
    res = engine.render("3b.docx", data, user_id="test_list")
    print(f"✅ Result: {res['pdf']}")

if __name__ == "__main__":
    create_template_3b()
    run_test()
