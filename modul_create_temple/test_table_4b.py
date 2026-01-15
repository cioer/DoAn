from docx import Document
from docx.shared import Pt
import os
import sys

# 1. TẠO TEMPLATE 4B (CHỈ CÓ HEADER)
def create_template_4b():
    doc = Document()
    doc.add_heading('MẪU 4B - DANH MỤC TỔNG HỢP', 0)
    
    p = doc.add_paragraph()
    p.add_run('HỘI ĐỒNG KHOA HỌC KHOA: ').bold = True
    p.add_run('{{ khoa }}') # Biến đơn
    
    doc.add_paragraph('\n')
    
    # Tạo bảng có 1 dòng header
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Tên đề tài'
    hdr_cells[2].text = 'Chủ nhiệm'
    hdr_cells[3].text = 'Mục tiêu'
    hdr_cells[4].text = 'Kinh phí'
    
    path = 'form_engine/templates/4b.docx'
    doc.save(path)
    print(f"Created template: {path}")

# 2. LOGIC ĐIỀN BẢNG (Dynamic Table Filling)
def run_table_test():
    # Dữ liệu mẫu (Giả sử lấy từ DB ra 5 dòng)
    context = {
        "khoa": "CƠ KHÍ ĐỘNG LỰC",
        "danh_sach": [
            {
                "stt": 1, 
                "ten": "Thiết kế máy bóc vỏ lạc tự động", 
                "cn": "TS. Trần Cơ Khí", 
                "mt": "Tăng năng suất 200%", 
                "kp": "15.000.000"
            },
            {
                "stt": 2, 
                "ten": "Nghiên cứu nhiên liệu sinh học từ rơm rạ", 
                "cn": "ThS. Lê Môi Trường", 
                "mt": "Giảm ô nhiễm", 
                "kp": "12.500.000"
            },
            {
                "stt": 3, 
                "ten": "Chế tạo Robot hàn tự hành", 
                "cn": "KS. Phạm Tự Động", 
                "mt": "Hàn chính xác 99%", 
                "kp": "30.000.000"
            },
            {
                "stt": 4, 
                "ten": "Hệ thống tưới tiêu thông minh IoT", 
                "cn": "TS. Nguyễn Nông Nghiệp", 
                "mt": "Tiết kiệm nước", 
                "kp": "8.000.000"
            },
             {
                "stt": 5, 
                "ten": "Xây dựng mô hình dạy học ảo VR", 
                "cn": "ThS. Đỗ Giáo Dục", 
                "mt": "Trực quan hóa", 
                "kp": "20.000.000"
            }
        ]
    }

    # Load Template
    template_path = 'form_engine/templates/4b.docx'
    doc = Document(template_path)
    
    # 1. Điền biến đơn (khoa)
    for p in doc.paragraphs:
        if "{{ khoa }}" in p.text:
            p.text = p.text.replace("{{ khoa }}", context['khoa'])
            
    # 2. Điền bảng (Tìm bảng đầu tiên để điền)
    table = doc.tables[0]
    
    # Loop qua danh sách data và thêm dòng mới
    for item in context['danh_sach']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['stt'])
        row_cells[1].text = item['ten']
        row_cells[2].text = item['cn']
        row_cells[3].text = item['mt']
        row_cells[4].text = item['kp']
    
    # Xuất file
    out_docx = 'form_engine/output/2026-01-15/4b_table_demo.docx'
    out_pdf = 'form_engine/output/2026-01-15/4b_table_demo.pdf'
    
    doc.save(out_docx)
    
    # Convert PDF
    import subprocess
    try:
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "form_engine/output/2026-01-15", out_docx]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"✅ Created Table Demo:\n   DOCX: {out_docx}\n   PDF:  {out_pdf}")
    except:
        print("Done (PDF failed)")

if __name__ == "__main__":
    create_template_4b()
    run_table_test()
