from docx import Document
import os

def verify_4b_result():
    # Tìm file 4b mới nhất
    path = "form_engine/output/2026-01-15/4b_221015.docx"
    if not os.path.exists(path):
        print("❌ File output not found.")
        return

    doc = Document(path)
    
    # Check Table Index 1
    if len(doc.tables) > 1:
        table = doc.tables[1]
        print(f"✅ Table found. Total rows: {len(table.rows)}")
        
        # Print Rows (Skip header rows if many)
        print("\n--- DATA ROWS CHECK ---")
        for i, row in enumerate(table.rows):
            # Chỉ in 2 dòng đầu (header) và 5 dòng cuối (data mới thêm)
            if i < 2 or i >= len(table.rows) - 5:
                cells = [c.text.strip().replace('\n', ' ')[:20] + "..." for c in row.cells]
                print(f"Row {i}: {cells}")
    else:
        print("❌ Table Index 1 missing in output file.")

if __name__ == "__main__":
    verify_4b_result()
