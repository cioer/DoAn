from docx import Document
import os

def verify_4b_detailed():
    path = "form_engine/output/2026-01-15/4b_formatted.docx"
    doc = Document(path)
    
    print("--- CHECKING HEADER ---")
    for p in doc.paragraphs[:10]: # Check 10 dòng đầu
        if "{{" in p.text:
            print(f"⚠️ MISSING VAR: {p.text}")
        else:
            print(f"OK: {p.text}")

    print("\n--- CHECKING TABLE ---")
    if len(doc.tables) > 1:
        table = doc.tables[1]
        for i, row in enumerate(table.rows):
            # Check từng ô
            for j, cell in enumerate(row.cells):
                if "{{" in cell.text:
                    print(f"⚠️ ROW {i} COL {j} MISSING: {cell.text}")
                if not cell.text.strip():
                    print(f"⚠️ ROW {i} COL {j} EMPTY")
    else:
        print("Table missing")

if __name__ == "__main__":
    verify_4b_detailed()
