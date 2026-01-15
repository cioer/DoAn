from docx import Document
import os
import glob

def verify_output():
    # Tìm file output mới nhất
    output_dir = "form_engine/output"
    files = glob.glob(f"{output_dir}/**/*.docx", recursive=True)
    files.sort(key=os.path.getmtime, reverse=True)
    
    if not files:
        print("❌ Không tìm thấy file output nào.")
        return

    # Lấy 2 file mới nhất (1b và 2b)
    target_files = files[:2]
    
    print(f"--- VERIFYING {len(target_files)} NEWEST FILES ---\n")

    for file_path in target_files:
        print(f"📄 Checking: {os.path.basename(file_path)}")
        try:
            doc = Document(file_path)
            full_text = []
            for p in doc.paragraphs:
                full_text.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            content = "\n".join(full_text)
            
            # 1. Check xem còn sót thẻ {{...}} không
            if "{{" in content and "}}" in content:
                print("   ⚠️  CẢNH BÁO: Vẫn còn thẻ {{...}} chưa được thay thế!")
            else:
                print("   ✅ SẠCH: Không còn thẻ template nào.")

            # 2. Check dữ liệu cụ thể
            if "Chatbot" in content:
                print("   ✅ DATA FOUND: Tìm thấy từ khóa 'Chatbot'")
            else:
                print("   ❌ DATA MISSING: Không thấy dữ liệu 'Chatbot'")
                
            if "25.000.000" in content or "Lê Thẩm Định" in content:
                 print("   ✅ DATA FOUND: Tìm thấy số tiền/tên người")

        except Exception as e:
            print(f"   ❌ Lỗi đọc file: {e}")
        print("-" * 30)

if __name__ == "__main__":
    verify_output()
